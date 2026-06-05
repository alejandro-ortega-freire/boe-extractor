import unittest
from datetime import date
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import RGBColor

from source.anexo_vi_writer import create_anexo_vi_docx
from source.models import BasicData, SummaryModule
from source.schedule import calculate_schedule


class AnexoVIWriterTests(unittest.TestCase):
    def test_anexo_vi_uses_anexo_iii_base_with_custom_titles(self):
        with TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "anexoVI.docx"
            create_anexo_vi_docx(
                BasicData(
                    codigo="TEST0101",
                    nombre="Certificado de prueba.",
                ),
                [SummaryModule(text="MF0001_1: Módulo de prueba (10 horas)")],
                "10h",
                output_path,
                schedule={},
                teacher_name="Docente",
            )

            doc = Document(output_path)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

        self.assertEqual(doc.sections[0].orientation, WD_ORIENT.LANDSCAPE)
        self.assertIn("ANEXO VI", text)
        self.assertIn("Informe de Evaluación Individualizado", text)
        self.assertIn("INFORME DE EVALUACIÓN INDIVIDUALIZADO", text)
        self.assertNotIn("ANEXO III", text)
        self.assertNotIn("Planificación didáctica", text)
        self.assertNotIn("PLANIFICACIÓN DIDÁCTICA DEL CURSO COMPLETO", text)

    def test_anexo_vi_table_matches_requested_mock_structure(self):
        with TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "anexoVI.docx"
            create_anexo_vi_docx(
                BasicData(codigo="TEST0101"),
                [
                    SummaryModule(
                        text="MF0001_1: Módulo con UFs (60 horas)",
                        ufs=[
                            "UF0001: Primera UF (30 horas)",
                            "UF0002: Segunda UF (30 horas)",
                        ],
                    ),
                    SummaryModule(text="MF0002_1: Módulo sin UFs (10 horas)"),
                    SummaryModule(text="MP0001: Módulo de prácticas (10 horas)"),
                ],
                "10h",
                output_path,
                schedule={},
                teacher_name="Docente",
            )

            table = Document(output_path).tables[0]

        self.assertEqual(len(table.rows), 8)
        self.assertEqual(len(table.columns), 11)
        self.assertIn("Nombre y apellidos del alumno", table.rows[0].cells[0].text)
        self.assertIn("Nombre del alumno Apellido 1 Apellido 2", table.rows[0].cells[2].text)
        self.assertIn("MÓDULOS FORMATIVOS", table.rows[1].cells[0].text)
        self.assertIn("EVALUACIÓN DURANTE EL PROCESO DE APRENDIZAJE", table.rows[1].cells[2].text)
        self.assertIn("PRUEBA DE EVALUACIÓN FINAL", table.rows[1].cells[6].text)
        self.assertIn("CALIFICACIÓN FINAL DEL", table.rows[1].cells[10].text)
        self.assertIn("E1", table.rows[3].cells[2].text)
        table_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        self.assertNotIn("Indicar código y denominación", table_text)
        self.assertNotIn("MP0001", table_text)
        self.assertIn("MF0001_1: Módulo con UFs", table.rows[4].cells[0].text)
        self.assertIn("UF0001: Primera UF", table.rows[4].cells[1].text)
        self.assertIn("UF0002: Segunda UF", table.rows[5].cells[1].text)
        self.assertIn("MF0002_1: Módulo sin UFs", table.rows[6].cells[0].text)
        self.assertIn("MF0002_1: Módulo sin UFs", table.rows[6].cells[1].text)
        self.assertIn("MF0002_1: Módulo sin UFs", table.rows[7].cells[0].text)
        self.assertIn("MF0002_1: Módulo sin UFs", table.rows[7].cells[1].text)
        self.assertIn("APTO", table.rows[4].cells[10].text)
        apto_runs = [
            run
            for paragraph in table.rows[4].cells[10].paragraphs
            for run in paragraph.runs
            if run.text
        ]
        self.assertTrue(apto_runs)
        self.assertTrue(all(run.bold for run in apto_runs))
        self.assertIn("PPF\n60%", table.rows[3].cells[7].text)
        self.assertIn("PPF\n60%", table.rows[3].cells[9].text)

        name_runs = [run for paragraph in table.rows[0].cells[2].paragraphs for run in paragraph.runs if run.text]
        self.assertTrue(name_runs)
        self.assertFalse(any(run.bold for run in name_runs))

        percent_runs = [
            run
            for paragraph in table.rows[1].cells[2].paragraphs + table.rows[1].cells[6].paragraphs
            for run in paragraph.runs
            if "calificación final" in run.text
        ]
        self.assertTrue(percent_runs)
        self.assertTrue(all(run.font.color.rgb == RGBColor(255, 255, 255) for run in percent_runs))
        self.assertIn('w:color="000000"', table._tbl.xml)
        self.assertNotIn('w:color="FFFFFF"', table._tbl.xml)

    def test_anexo_vi_places_scores_only_in_uf_activity_columns(self):
        modules = [
            SummaryModule(
                text="MF0001_1: Módulo con UFs (60 horas)",
                ufs=[
                    "UF0001: Primera UF (30 horas)",
                    "UF0002: Segunda UF (30 horas)",
                ],
            )
        ]
        schedule = calculate_schedule(modules, 6, date(2026, 6, 1), custom_holidays={})

        with TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "anexoVI.docx"
            create_anexo_vi_docx(
                BasicData(codigo="TEST0101"),
                modules,
                "60h",
                output_path,
                schedule=schedule,
                teacher_name="Docente",
            )

            table = Document(output_path).tables[0]

        self.assertEqual(table.rows[4].cells[2].text, "Nota")
        self.assertEqual(table.rows[4].cells[3].text, "Nota")
        self.assertEqual(table.rows[4].cells[4].text, "")
        self.assertIn("F2F2F2", table.rows[4].cells[4]._tc.xml)
        self.assertEqual(table.rows[5].cells[2].text, "")
        self.assertEqual(table.rows[5].cells[3].text, "")
        self.assertEqual(table.rows[5].cells[4].text, "Nota")
        self.assertIn("F2F2F2", table.rows[5].cells[2]._tc.xml)
        self.assertIn("F2F2F2", table.rows[5].cells[3]._tc.xml)
        self.assertIn("ΣEvaluables / 3", table.rows[4].cells[5].text)
        self.assertEqual(table.rows[4].cells[10].text, table.rows[5].cells[10].text)
        self.assertEqual(table.rows[5].cells[6].text, "Nota final")
        self.assertEqual(table.rows[4].cells[8].text, "")
        self.assertEqual(table.rows[4].cells[9].text, "")
        self.assertEqual(table.rows[5].cells[8].text, "")
        self.assertIn("F2F2F2", table.rows[4].cells[8]._tc.xml)
        self.assertIn("F2F2F2", table.rows[4].cells[9]._tc.xml)
        self.assertIn("F2F2F2", table.rows[5].cells[8]._tc.xml)

        score_runs = [
            run
            for row in table.rows[4:6]
            for cell in row.cells[2:5]
            for paragraph in cell.paragraphs
            for run in paragraph.runs
            if run.text == "Nota"
        ]
        self.assertTrue(score_runs)
        self.assertTrue(all(not run.bold for run in score_runs))
        self.assertTrue(all(run.font.color.rgb == RGBColor(255, 0, 0) for run in score_runs))

        summary_runs = [
            run
            for cell in (table.rows[5].cells[6],)
            for paragraph in cell.paragraphs
            for run in paragraph.runs
            if run.text == "Nota final"
        ]
        self.assertEqual(len(summary_runs), 1)
        self.assertTrue(all(run.bold for run in summary_runs))
        self.assertTrue(all(run.font.size.pt == 10 for run in summary_runs))

    def test_anexo_vi_module_without_ufs_has_regular_and_summary_final_scores(self):
        modules = [SummaryModule(text="MF0001_1: Módulo sin UFs (30 horas)")]
        schedule = calculate_schedule(modules, 6, date(2026, 6, 1), custom_holidays={})

        with TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "anexoVI.docx"
            create_anexo_vi_docx(
                BasicData(codigo="TEST0101"),
                modules,
                "30h",
                output_path,
                schedule=schedule,
                teacher_name="Docente",
            )

            table = Document(output_path).tables[0]

        self.assertEqual(len(table.rows), 6)
        self.assertEqual(table.rows[4].height_rule, WD_ROW_HEIGHT_RULE.EXACTLY)
        self.assertEqual(table.rows[4].cells[4].text, "Nota")
        self.assertEqual(table.rows[5].cells[2].text, table.rows[4].cells[2].text)
        self.assertEqual(table.rows[4].cells[5].text, "Nota")
        self.assertEqual(table.rows[4].cells[6].text, "")
        self.assertEqual(table.rows[4].cells[7].text, "")
        self.assertEqual(table.rows[5].cells[4].text, "Nota final")
        self.assertEqual(table.rows[5].cells[6].text, "")

        summary_runs = [
            run
            for cell in (table.rows[5].cells[4],)
            for paragraph in cell.paragraphs
            for run in paragraph.runs
            if run.text == "Nota final"
        ]
        self.assertEqual(len(summary_runs), 1)
        self.assertTrue(all(run.bold for run in summary_runs))

    def test_anexo_vi_average_formula_uses_each_module_evaluable_count(self):
        modules = [
            SummaryModule(text="MF0001_1: Módulo corto (30 horas)"),
            SummaryModule(text="MF0002_1: Módulo largo (60 horas)"),
        ]
        schedule = calculate_schedule(modules, 6, date(2026, 6, 1), custom_holidays={})

        with TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "anexoVI.docx"
            create_anexo_vi_docx(
                BasicData(codigo="TEST0101"),
                modules,
                "90h",
                output_path,
                schedule=schedule,
                teacher_name="Docente",
            )

            table = Document(output_path).tables[0]

        self.assertIn("ΣEvaluables / 1", table.rows[4].cells[5].text)
        self.assertIn("ΣEvaluables / 3", table.rows[6].cells[5].text)


if __name__ == "__main__":
    unittest.main()
