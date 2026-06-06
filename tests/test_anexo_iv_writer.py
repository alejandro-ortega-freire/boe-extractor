import unittest
from datetime import date
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.shared import RGBColor

from source.anexo_iv_writer import create_anexo_iv_docx
from source.models import BasicData, Criterion, SummaryModule, TrainingModule
from source.schedule import calculate_schedule


class AnexoIVWriterTests(unittest.TestCase):
    def test_anexo_iv_adds_notes_after_table(self):
        module = TrainingModule(
            identifier="MF0001_1: Módulo de prueba",
            hours="20",
            objective="Objetivo de prueba.",
            criteria=[Criterion(text="C1: Primer criterio.")],
        )

        with TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "anexoIV.docx"
            create_anexo_iv_docx(
                BasicData(codigo="TEST0101"),
                module,
                "20 horas",
                output_path,
                schedule={},
                add_header_footer=lambda doc, teacher_name: None,
            )

            text = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)

        self.assertIn("1  Incluir los Resultados de Aprendizaje (RA)", text)
        self.assertIn("2  Introducir los contenidos", text)
        self.assertIn("3  Especificar las diferentes acciones", text)
        self.assertIn("4  Indicar los que corresponden exclusivamente", text)

    def test_strategy_lorem_placeholders_are_red_only(self):
        module = TrainingModule(
            identifier="MF0001_1: Módulo de prueba",
            hours="20",
            objective="Objetivo de prueba.",
            criteria=[Criterion(text="C1: Primer criterio.")],
        )

        with TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "anexoIV.docx"
            create_anexo_iv_docx(
                BasicData(codigo="TEST0101"),
                module,
                "20 horas",
                output_path,
                schedule={},
                add_header_footer=lambda doc, teacher_name: None,
            )

            strategy_cell = Document(output_path).tables[0].rows[1].cells[2]

        lorem_runs = [
            run
            for paragraph in strategy_cell.paragraphs
            for run in paragraph.runs
            if run.text == "Lorem ipsum dolor met"
        ]
        title_run = next(
            run
            for paragraph in strategy_cell.paragraphs
            for run in paragraph.runs
            if run.text == "SESIÓN #"
        )

        self.assertTrue(lorem_runs)
        self.assertTrue(all(run.font.color.rgb == RGBColor(192, 0, 0) for run in lorem_runs))
        self.assertNotEqual(title_run.font.color.rgb, RGBColor(192, 0, 0))

    def test_first_strategy_cell_includes_evaluation_suggestion(self):
        module = TrainingModule(
            identifier="MF0001_1: Módulo de prueba",
            hours="60",
            objective="Objetivo de prueba.",
            criteria=[
                Criterion(text="C1: Primer criterio."),
                Criterion(text="C2: Segundo criterio."),
            ],
        )
        schedule = calculate_schedule(
            [SummaryModule(text="MF0001_1: Módulo de prueba (60 horas)")],
            6,
            date(2026, 6, 1),
            custom_holidays={},
        )

        with TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "anexoIV.docx"
            create_anexo_iv_docx(
                BasicData(codigo="TEST0101"),
                module,
                "60 horas",
                output_path,
                schedule=schedule,
                add_header_footer=lambda doc, teacher_name: None,
            )

            table = Document(output_path).tables[0]

        first_strategy_cell = table.rows[1].cells[2]
        second_strategy_cell = table.rows[2].cells[2]

        self.assertIn("SESIÓN #", first_strategy_cell.text)
        self.assertIn("Sugerencia:", first_strategy_cell.text)
        self.assertIn("SESIÓN 1 ⭢ Presentación del MF", first_strategy_cell.text)
        self.assertIn("SESIÓN 3 ⭢ E1", first_strategy_cell.text)
        self.assertIn("SESIÓN 9 ⭢ Prueba Final", first_strategy_cell.text)
        self.assertIn("SESIÓN 10 ⭢ Recuperación", first_strategy_cell.text)
        self.assertNotIn("Sugerencia:", second_strategy_cell.text)

        suggestion_run = next(
            run
            for paragraph in first_strategy_cell.paragraphs
            for run in paragraph.runs
            if run.text == "Sugerencia:"
        )
        event_run = next(
            run
            for paragraph in first_strategy_cell.paragraphs
            for run in paragraph.runs
            if run.text == "SESIÓN 3 ⭢ E1"
        )

        self.assertTrue(suggestion_run.bold)
        self.assertEqual(suggestion_run.font.color.rgb, RGBColor(192, 0, 0))
        self.assertTrue(event_run.italic)
        self.assertEqual(event_run.font.color.rgb, RGBColor(192, 0, 0))


if __name__ == "__main__":
    unittest.main()
