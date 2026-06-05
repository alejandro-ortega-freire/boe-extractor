import unittest
from datetime import date
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import RGBColor

from source.anexo_v_writer import (
    EVALUABLE_ACTIVITY_DESCRIPTION,
    FINAL_EVALUATION_HEADER,
    FINAL_EVALUATION_ITEMS,
    RECOVERY_EVALUATION_HEADER,
    RECOVERY_EVALUATION_INTRO,
    RECOVERY_EVALUATION_ITEMS,
    build_anexo_v_filename,
    clean_evaluation_space_name,
    create_anexo_v_docx,
)
from source.models import BasicData, SummaryModule, TrainingModule, TrainingUnit
from source.schedule import calculate_schedule


class AnexoVWriterTests(unittest.TestCase):
    def test_anexo_v_uses_evaluation_planning_header_texts(self):
        with TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "anexoV.docx"
            create_anexo_v_docx(
                BasicData(
                    codigo="TEST0101",
                    nombre="Certificado de prueba.",
                    familia="Familia de prueba.",
                    nivel="2.",
                ),
                TrainingModule(
                    identifier="MF0001_1: Módulo de prueba",
                    hours="100",
                    objective="Objetivo de prueba.",
                ),
                "100 horas",
                output_path,
                schedule={},
                add_header_footer=lambda doc, teacher_name: None,
                teacher_name="Docente",
            )

            text = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)

        self.assertIn("ANEXO V - MF0001_1 Módulo de prueba", text)
        self.assertIn("Planificación de la evaluación del aprendizaje", text)
        self.assertIn("PLANIFICACIÓN DE LA EVALUACIÓN DEL APRENDIZAJE", text)
        self.assertIn("1 Identificar las actividades e instrumentos de evaluación", text)
        self.assertIn("2 Las fechas de evaluación estarán actualizadas", text)
        self.assertNotIn("ANEXO IV", text)
        self.assertNotIn("PROGRAMACIÓN DIDÁCTICA DEL MÓDULO PROFESIONAL", text)

    def test_anexo_v_filename_matches_module_and_certificate(self):
        self.assertEqual(
            build_anexo_v_filename("MF0001_1", "TEST0101"),
            "anexoV_MF0001_1_TEST0101.docx",
        )

    def test_anexo_v_table_adds_one_row_per_uf(self):
        with TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "anexoV.docx"
            create_anexo_v_docx(
                BasicData(codigo="TEST0101"),
                TrainingModule(
                    identifier="MF0001_1: Módulo de prueba",
                    ufs=[
                        TrainingUnit(code="UF0001", name="Primera unidad"),
                        TrainingUnit(code="UF0002", name="Segunda unidad"),
                    ],
                ),
                "100 horas",
                output_path,
                schedule={},
                add_header_footer=lambda doc, teacher_name: None,
            )

            table = Document(output_path).tables[0]

        self.assertEqual(len(table.rows), 4)
        self.assertEqual(table.rows[0].cells[0].text, "MÓDULO\nPROFESIONAL")
        self.assertEqual(table.rows[1].cells[0].text, "BLOQUES\nFORMATIVOS")
        self.assertEqual(table.rows[0].cells[1].text, "DURANTE EL PROCESO DE\nAPRENDIZAJE")
        self.assertEqual(table.rows[1].cells[1].text, "ACTIVIDADES E INSTRUMENTOS\nDE EVALUACIÓN¹")
        self.assertEqual(table.rows[0].cells[2].text, "Realización de la evaluación")
        self.assertEqual(table.rows[1].cells[2].text, "Espacios")
        self.assertEqual(table.rows[1].cells[3].text, "Duración")
        self.assertEqual(table.rows[1].cells[4].text, "Fechas de\nevaluación²")
        self.assertEqual(table.rows[2].cells[0].text, "UF0001: Primera unidad")
        self.assertEqual(table.rows[3].cells[0].text, "UF0002: Segunda unidad")

    def test_anexo_v_table_uses_module_name_when_module_has_no_ufs(self):
        with TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "anexoV.docx"
            create_anexo_v_docx(
                BasicData(codigo="TEST0101"),
                TrainingModule(identifier="MF0001_1: Módulo de prueba"),
                "100 horas",
                output_path,
                schedule={},
                add_header_footer=lambda doc, teacher_name: None,
            )

            table = Document(output_path).tables[0]

        self.assertEqual(len(table.rows), 3)
        self.assertEqual(table.rows[2].cells[0].text, "MF0001_1: Módulo de prueba")

    def test_anexo_v_splits_module_row_into_evaluation_events(self):
        schedule = calculate_schedule(
            [SummaryModule(text="MF0001_1: Módulo de prueba (60 horas)")],
            6,
            date(2026, 6, 1),
            custom_holidays={},
        )

        with TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "anexoV.docx"
            create_anexo_v_docx(
                BasicData(codigo="TEST0101"),
                TrainingModule(
                    identifier="MF0001_1: Módulo de prueba",
                    hours="60",
                ),
                "60 horas",
                output_path,
                schedule=schedule,
                add_header_footer=lambda doc, teacher_name: None,
            )

            table = Document(output_path).tables[0]
            labels = [row.cells[1].text.splitlines()[0] for row in table.rows[2:]]
            durations = [row.cells[3].text for row in table.rows[2:]]
            dates = [row.cells[4].text for row in table.rows[2:]]

        self.assertEqual(
            labels,
            [
                "E1: Actividad Evaluable",
                "E2: Actividad Evaluable",
                "E3: Actividad Evaluable",
                "PRUEBA DE EVALUACIÓN FINAL",
                "1. Parte Teórica: Prueba teórica sobre los contenidos del manual y los contenidos específicos dados en las sesiones de formación.",
                "PRUEBA DE RECUPERACIÓN DE",
                "Re-evaluación de la Parte Teórica (POF) y",
            ],
        )
        self.assertEqual(durations, ["4 horas", "4 horas", "4 horas", "Duración", "6 horas", "Duración", "6 horas"])
        self.assertIn("(Sesión 9)", dates[-3])
        self.assertIn("(Sesión 10)", dates[-1])

    def test_anexo_v_session_numbers_restart_for_each_module(self):
        schedule = calculate_schedule(
            [
                SummaryModule(text="MF0001_1: Primer módulo (30 horas)"),
                SummaryModule(text="MF0002_1: Segundo módulo (30 horas)"),
            ],
            6,
            date(2026, 6, 1),
            custom_holidays={},
        )

        with TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "anexoV.docx"
            create_anexo_v_docx(
                BasicData(codigo="TEST0101"),
                TrainingModule(
                    identifier="MF0002_1: Segundo módulo",
                    hours="30",
                ),
                "60 horas",
                output_path,
                schedule=schedule,
                add_header_footer=lambda doc, teacher_name: None,
            )

            table = Document(output_path).tables[0]
            dates = [
                row.cells[4].text
                for row in table.rows[2:]
                if "Sesión" in row.cells[4].text
            ]

        self.assertIn("(Sesión 4)", dates[-2])
        self.assertIn("(Sesión 5)", dates[-1])
        self.assertNotIn("(Sesión 9)", dates[-2])
        self.assertNotIn("(Sesión 10)", dates[-1])

    def test_anexo_v_uses_clean_red_training_spaces(self):
        schedule = calculate_schedule(
            [SummaryModule(text="MF0001_1: Módulo de prueba (30 horas)")],
            6,
            date(2026, 6, 1),
            custom_holidays={},
        )

        with TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "anexoV.docx"
            create_anexo_v_docx(
                BasicData(codigo="TEST0101"),
                TrainingModule(
                    identifier="MF0001_1: Módulo de prueba",
                    hours="30",
                ),
                "30 horas",
                output_path,
                schedule=schedule,
                spaces=[
                    "Aula de gestión de 45 m2 (para 15 alumnos) o de 60 m2 (para 25 alumnos)",
                    "Taller de prácticas de 30 m2 (para 15 alumnos) o de 50 m2 (para 25 alumnos)",
                ],
                add_header_footer=lambda doc, teacher_name: None,
            )

            cell = Document(output_path).tables[0].rows[2].cells[2]
            runs = [
                run
                for paragraph in cell.paragraphs
                for run in paragraph.runs
            ]

        self.assertEqual(cell.text, "Aula de gestión\n\nTaller de prácticas")
        self.assertTrue(runs)
        self.assertTrue(all(run.font.color.rgb == RGBColor(192, 0, 0) for run in runs))

    def test_anexo_v_uses_black_training_space_when_there_is_only_one(self):
        schedule = calculate_schedule(
            [SummaryModule(text="MF0001_1: Módulo de prueba (30 horas)")],
            6,
            date(2026, 6, 1),
            custom_holidays={},
        )

        with TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "anexoV.docx"
            create_anexo_v_docx(
                BasicData(codigo="TEST0101"),
                TrainingModule(
                    identifier="MF0001_1: Módulo de prueba",
                    hours="30",
                ),
                "30 horas",
                output_path,
                schedule=schedule,
                spaces=[
                    "Aula de gestión de 45 m2 (para 15 alumnos) o de 60 m2 (para 25 alumnos)",
                ],
                add_header_footer=lambda doc, teacher_name: None,
            )

            cell = Document(output_path).tables[0].rows[2].cells[2]
            runs = [
                run
                for paragraph in cell.paragraphs
                for run in paragraph.runs
                if run.text
            ]

        self.assertEqual(cell.text, "Aula de gestión")
        self.assertTrue(runs)
        self.assertTrue(all(run.font.color.rgb is None for run in runs))

    def test_anexo_v_marks_only_activity_durations_in_red(self):
        schedule = calculate_schedule(
            [SummaryModule(text="MF0001_1: Módulo de prueba (60 horas)")],
            6,
            date(2026, 6, 1),
            custom_holidays={},
        )

        with TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "anexoV.docx"
            create_anexo_v_docx(
                BasicData(codigo="TEST0101"),
                TrainingModule(
                    identifier="MF0001_1: Módulo de prueba",
                    hours="60",
                ),
                "60 horas",
                output_path,
                schedule=schedule,
                add_header_footer=lambda doc, teacher_name: None,
            )

            table = Document(output_path).tables[0]
            activity_duration = table.rows[2].cells[3].paragraphs[0].runs[0]
            final_duration = table.rows[-2].cells[3].paragraphs[0].runs[0]
            recovery_duration = table.rows[-1].cells[3].paragraphs[0].runs[0]

        self.assertEqual(activity_duration.font.color.rgb, RGBColor(192, 0, 0))
        self.assertIsNone(final_duration.font.color.rgb)
        self.assertIsNone(recovery_duration.font.color.rgb)

    def test_anexo_v_data_cells_are_top_aligned(self):
        schedule = calculate_schedule(
            [SummaryModule(text="MF0001_1: Módulo de prueba (60 horas)")],
            6,
            date(2026, 6, 1),
            custom_holidays={},
        )

        with TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "anexoV.docx"
            create_anexo_v_docx(
                BasicData(codigo="TEST0101"),
                TrainingModule(
                    identifier="MF0001_1: Módulo de prueba",
                    hours="60",
                ),
                "60 horas",
                output_path,
                schedule=schedule,
                add_header_footer=lambda doc, teacher_name: None,
            )

            first_data_row = Document(output_path).tables[0].rows[2]

        self.assertEqual(first_data_row.cells[0].vertical_alignment, WD_ALIGN_VERTICAL.TOP)
        self.assertEqual(first_data_row.cells[1].vertical_alignment, WD_ALIGN_VERTICAL.TOP)
        self.assertEqual(first_data_row.cells[2].vertical_alignment, WD_ALIGN_VERTICAL.CENTER)
        self.assertEqual(first_data_row.cells[3].vertical_alignment, WD_ALIGN_VERTICAL.CENTER)
        self.assertEqual(first_data_row.cells[4].vertical_alignment, WD_ALIGN_VERTICAL.CENTER)

    def test_anexo_v_activity_cell_has_bold_title_and_red_description(self):
        schedule = calculate_schedule(
            [SummaryModule(text="MF0001_1: Módulo de prueba (30 horas)")],
            6,
            date(2026, 6, 1),
            custom_holidays={},
        )

        with TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "anexoV.docx"
            create_anexo_v_docx(
                BasicData(codigo="TEST0101"),
                TrainingModule(
                    identifier="MF0001_1: Módulo de prueba",
                    hours="30",
                ),
                "30 horas",
                output_path,
                schedule=schedule,
                add_header_footer=lambda doc, teacher_name: None,
            )

            cell = Document(output_path).tables[0].rows[2].cells[1]
            title_run = next(run for run in cell.paragraphs[0].runs if run.text)
            description_run = next(run for run in cell.paragraphs[1].runs if run.text)

        self.assertEqual(title_run.text, "E1: Actividad Evaluable")
        self.assertTrue(title_run.bold)
        self.assertEqual(description_run.text, EVALUABLE_ACTIVITY_DESCRIPTION)
        self.assertEqual(description_run.font.color.rgb, RGBColor(192, 0, 0))

    def test_anexo_v_final_evaluation_uses_header_and_red_theory_practice_content(self):
        schedule = calculate_schedule(
            [SummaryModule(text="MF0001_1: Módulo de prueba (60 horas)")],
            6,
            date(2026, 6, 1),
            custom_holidays={},
        )

        with TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "anexoV.docx"
            create_anexo_v_docx(
                BasicData(codigo="TEST0101"),
                TrainingModule(
                    identifier="MF0001_1: Módulo de prueba",
                    hours="60",
                ),
                "60 horas",
                output_path,
                schedule=schedule,
                add_header_footer=lambda doc, teacher_name: None,
            )

            table = Document(output_path).tables[0]
            header_index = next(
                index
                for index, row in enumerate(table.rows)
                if row.cells[1].text == FINAL_EVALUATION_HEADER
            )
            header_cell = table.rows[header_index].cells[1]
            content_cell = table.rows[header_index + 1].cells[1]

        self.assertEqual(header_cell.text, FINAL_EVALUATION_HEADER)
        self.assertEqual(table.rows[header_index].cells[4].text, "Fechas de\nevaluación²")
        self.assertEqual(
            content_cell.text,
            "".join(FINAL_EVALUATION_ITEMS[0]) + "\n" + "".join(FINAL_EVALUATION_ITEMS[1]),
        )

        content_runs = [
            run
            for paragraph in content_cell.paragraphs
            for run in paragraph.runs
            if run.text
        ]

        self.assertTrue(all(run.font.color.rgb == RGBColor(192, 0, 0) for run in content_runs))
        self.assertTrue(content_runs[0].bold)
        self.assertFalse(content_runs[1].bold)
        self.assertTrue(content_runs[2].bold)
        self.assertFalse(content_runs[3].bold)

    def test_anexo_v_recovery_evaluation_uses_header_and_red_theory_practice_content(self):
        schedule = calculate_schedule(
            [SummaryModule(text="MF0001_1: Módulo de prueba (60 horas)")],
            6,
            date(2026, 6, 1),
            custom_holidays={},
        )

        with TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "anexoV.docx"
            create_anexo_v_docx(
                BasicData(codigo="TEST0101"),
                TrainingModule(
                    identifier="MF0001_1: Módulo de prueba",
                    hours="60",
                ),
                "60 horas",
                output_path,
                schedule=schedule,
                add_header_footer=lambda doc, teacher_name: None,
            )

            table = Document(output_path).tables[0]
            header_index = next(
                index
                for index, row in enumerate(table.rows)
                if row.cells[1].text == RECOVERY_EVALUATION_HEADER
            )
            header_cell = table.rows[header_index].cells[1]
            content_cell = table.rows[header_index + 1].cells[1]

        self.assertEqual(header_cell.text, RECOVERY_EVALUATION_HEADER)
        self.assertEqual(table.rows[header_index].cells[4].text, "Fechas de\nevaluación²")
        self.assertEqual(
            content_cell.text,
            (
                RECOVERY_EVALUATION_INTRO
                + "\n"
                + "".join(RECOVERY_EVALUATION_ITEMS[0])
                + "\n"
                + "".join(RECOVERY_EVALUATION_ITEMS[1])
            ),
        )

        content_runs = [
            run
            for paragraph in content_cell.paragraphs
            for run in paragraph.runs
            if run.text
        ]

        self.assertTrue(all(run.font.color.rgb == RGBColor(192, 0, 0) for run in content_runs))
        self.assertFalse(content_runs[0].bold)
        self.assertTrue(content_runs[1].bold)
        self.assertFalse(content_runs[2].bold)
        self.assertTrue(content_runs[3].bold)
        self.assertFalse(content_runs[4].bold)

    def test_clean_evaluation_space_name_removes_area_and_capacity(self):
        self.assertEqual(
            clean_evaluation_space_name(
                "Aula de gestión de 45 m2 (para 15 alumnos) o de 60 m2 (para 25 alumnos)"
            ),
            "Aula de gestión",
        )


if __name__ == "__main__":
    unittest.main()
