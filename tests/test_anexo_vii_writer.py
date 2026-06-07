import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import RGBColor

from source.anexo_vii_writer import create_anexo_vii_docx
from source.docx_styles import ANEXO_III_HEADER_FILL
from source.models import BasicData, SummaryModule


class AnexoVIIWriterTests(unittest.TestCase):
    def test_anexo_vii_uses_anexo_vi_base_with_custom_titles(self):
        with TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "anexoVII.docx"
            create_anexo_vii_docx(
                BasicData(
                    codigo="TEST0101",
                    nombre="Certificado de prueba.",
                ),
                [
                    SummaryModule(
                        text="MF0001_1: Módulo con UFs (60 horas)",
                        ufs=[
                            "UF0001: Primera UF (30 horas)",
                            "UF0002: Segunda UF (30 horas)",
                        ],
                    ),
                    SummaryModule(text="MF0002_1: Módulo sin UFs (10 horas)"),
                    SummaryModule(text="MP0001: Módulo de prácticas (10 horas)"),
                ],
                "10h",
                output_path,
                schedule={},
                teacher_name="Docente",
                student_count=3,
            )

            doc = Document(output_path)
            paragraph_texts = [paragraph.text for paragraph in doc.paragraphs]
            text = "\n".join(paragraph_texts)
            table = doc.tables[0]
            signature_table = doc.tables[1]
            table_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
            signature_table_text = "\n".join(
                cell.text
                for row in signature_table.rows
                for cell in row.cells
            )

        self.assertEqual(doc.sections[0].orientation, WD_ORIENT.LANDSCAPE)
        self.assertEqual(len(doc.tables), 2)
        self.assertIn("ANEXO VII", text)
        self.assertIn("Acta de Evaluación", text)
        self.assertIn("ACTA DE EVALUACIÓN", text)
        self.assertEqual(len(table.rows), 5)
        self.assertIn("Nº", table.rows[0].cells[0].text)
        self.assertIn("DNI/NIE", table.rows[0].cells[1].text)
        self.assertIn("APELLIDOS/NOMBRE", table.rows[0].cells[2].text)
        self.assertIn("MF 1\n(MF0001_1)", table.rows[0].cells[3].text)
        self.assertIn("UF 1\n(UF0001)", table.rows[1].cells[3].text)
        self.assertIn("UF 2\n(UF0002)", table.rows[1].cells[4].text)
        self.assertIn("CALIFICACIÓN", table.rows[1].cells[5].text)
        self.assertIn("MF 2\n(MF0002_1)", table_text)
        self.assertNotIn("En caso de no", table_text)
        self.assertEqual(table.rows[0].cells[-3].text, "MP\n(MP0001)")
        self.assertIn("PROPUESTA\nCERTIFICADO", table.rows[0].cells[-2].text)
        self.assertIn("PROPUESTA\nACREDITACIÓN\nPARCIAL", table.rows[0].cells[-1].text)
        self.assertEqual(table.rows[2].cells[0].text, "1")
        self.assertEqual(table.rows[3].cells[0].text, "2")
        self.assertEqual(table.rows[4].cells[0].text, "3")
        self.assertNotEqual(table.rows[2].height_rule, WD_ROW_HEIGHT_RULE.EXACTLY)
        self.assertEqual(table.rows[3].height_rule, WD_ROW_HEIGHT_RULE.EXACTLY)
        self.assertEqual(table.rows[4].height_rule, WD_ROW_HEIGHT_RULE.EXACTLY)
        self.assertIn("00000000-L", table.rows[2].cells[1].text)
        self.assertIn("Apellido 1 Apellido 2, Nombre", table.rows[2].cells[2].text)
        name_runs = [
            run
            for paragraph in table.rows[2].cells[2].paragraphs
            for run in paragraph.runs
            if run.text.strip()
        ]
        self.assertTrue(name_runs)
        self.assertTrue(all(run.bold for run in name_runs))
        self.assertTrue(all(not cell.text for cell in table.rows[3].cells[1:]))
        self.assertTrue(all(not cell.text for cell in table.rows[4].cells[1:]))
        self.assertIn("Apto /No apto", table_text)
        self.assertIn("Apto (suficiente)/", table_text)
        self.assertIn("Apto (notable)/", table_text)
        self.assertIn("Apto (sobresaliente)", table_text)
        self.assertIn("/Exento", table_text)
        self.assertIn("SI/NO", table_text)
        self.assertIn(f'w:fill="{ANEXO_III_HEADER_FILL}"', table.rows[0].cells[0]._tc.xml)
        header_runs = [
            run
            for row in table.rows[:2]
            for cell in row.cells
            for paragraph in cell.paragraphs
            for run in paragraph.runs
            if run.text.strip()
        ]
        body_runs = [
            run
            for row in table.rows[2:]
            for cell in row.cells
            for paragraph in cell.paragraphs
            for run in paragraph.runs
            if run.text.strip()
        ]
        self.assertTrue(header_runs)
        self.assertTrue(all(run.font.color.rgb == RGBColor(255, 255, 255) for run in header_runs))
        self.assertTrue(body_runs)
        self.assertTrue(all(run.font.color.rgb == RGBColor(255, 0, 0) for run in body_runs))
        self.assertNotIn("ANEXO VI", paragraph_texts)
        self.assertNotIn("Informe de Evaluación Individualizado", text)
        self.assertNotIn("INFORME DE EVALUACIÓN INDIVIDUALIZADO", text)
        self.assertEqual(len(signature_table.columns), 4)
        self.assertIn("MF1\n(MF0001_1: Módulo con UFs)", signature_table.rows[0].cells[0].text)
        self.assertIn("MF2\n(MF0002_1: Módulo sin UFs)", signature_table.rows[0].cells[1].text)
        self.assertIn("Formador/a:", signature_table_text)
        self.assertIn("Docente", signature_table_text)
        self.assertIn("Firma:", signature_table_text)
        self.assertEqual("", signature_table.rows[0].cells[-2].text)
        self.assertIn("Responsable/Dirección", signature_table.rows[0].cells[-1].text)
        self.assertIn("Firma:", signature_table.rows[0].cells[-1].text)
        self.assertNotIn("FIRMA DEL/LA DOCENTE QUE", signature_table_text)
        self.assertNotIn("MP0001", signature_table_text)
        self.assertEqual(signature_table.rows[0].height_rule, WD_ROW_HEIGHT_RULE.AT_LEAST)
        signature_runs = [
            run
            for paragraph in signature_table.rows[0].cells[0].paragraphs
            for run in paragraph.runs
            if run.text.strip()
        ]
        bold_signature_texts = {run.text for run in signature_runs if run.bold}
        self.assertIn("MF1", bold_signature_texts)
        self.assertIn("Formador/a:", bold_signature_texts)
        self.assertIn("Firma:", bold_signature_texts)
        self.assertNotIn("Docente", bold_signature_texts)
        responsible_runs = [
            run
            for paragraph in signature_table.rows[0].cells[-1].paragraphs
            for run in paragraph.runs
            if run.text.strip()
        ]
        bold_responsible_texts = {run.text for run in responsible_runs if run.bold}
        self.assertIn("Responsable/Dirección", bold_responsible_texts)
        self.assertIn("Firma:", bold_responsible_texts)


if __name__ == "__main__":
    unittest.main()
