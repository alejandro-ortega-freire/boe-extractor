import os

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from source.docx_styles import (
    DOC_HEADER_LOGO_WIDTH_INCHES,
    DOC_HEADER_SPACE_AFTER_PT,
    DOC_HEADER_TABLE_WIDTH_INCHES,
    DOC_HEADER_TEACHER_FONT_SIZE,
    DOC_HEADER_TOP_MARGIN_INCHES,
)
from source.normalization import normalize_text
from source.settings import DEFAULT_TEACHER_NAME, LOGO_PATH


def safe_text(text):
    return normalize_text(text)


def add_page_number(paragraph):
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def clear_paragraph(paragraph):
    paragraph._p.clear_content()


def add_horizontal_rule(doc, *, color="000000", space_before=12, space_after=12, size=8, space=1):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)

    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    return paragraph


def add_header_footer(doc, teacher_name=DEFAULT_TEACHER_NAME):
    for section in doc.sections:
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        section.top_margin = Inches(DOC_HEADER_TOP_MARGIN_INCHES)

        header = section.header
        header_table = header.add_table(rows=1, cols=2, width=Inches(DOC_HEADER_TABLE_WIDTH_INCHES))
        header_table.autofit = True

        left = header_table.cell(0, 0)
        right = header_table.cell(0, 1)
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        left_p = left.paragraphs[0]
        left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_p.paragraph_format.space_after = Pt(DOC_HEADER_SPACE_AFTER_PT)

        if os.path.exists(LOGO_PATH):
            left_p.add_run().add_picture(LOGO_PATH, width=Inches(DOC_HEADER_LOGO_WIDTH_INCHES))
        else:
            run = left_p.add_run("BOExtractor")
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 48, 112)

        right_p = right.paragraphs[0]
        right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_p.paragraph_format.space_after = Pt(DOC_HEADER_SPACE_AFTER_PT)
        run = right_p.add_run(safe_text(teacher_name))
        run.font.size = Pt(DOC_HEADER_TEACHER_FONT_SIZE)
        run.bold = True

        footer = section.footer
        paragraph = footer.paragraphs[0]
        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_page_number(paragraph)
