from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
import re

from source.anexo_iv_writer import (
    add_horizontal_line,
    add_module_header,
    configure_page,
    module_identifier_without_hours,
)
from source.docx_styles import (
    ANEXO_IV_FONT_SIZE,
    ANEXO_IV_TABLE_HEADER_FILL,
    ANEXO_IV_TABLE_WIDTH_PERCENT,
    LIGHT_BORDER,
)
from source.evaluation_plan import (
    block_code,
    build_evaluation_events,
    duration_text,
    evaluable_activity_count,
    evaluation_blocks,
    evaluation_date_text,
)
from source.settings import DEFAULT_TEACHER_NAME
from source.table_styles import (
    apply_vertical_borders,
    set_cell_shading,
    set_cell_text,
    set_table_width_percent,
)


EVALUABLE_ACTIVITY_DESCRIPTION = (
    "Desarrollo de una actividad práctica orientada a la aplicación de los "
    "contenidos del módulo formativo. El alumnado deberá realizar tareas "
    "relacionadas con los resultados de aprendizaje establecidos, aplicando "
    "los procedimientos y herramientas correspondientes, y documentando el "
    "trabajo realizado."
)
FINAL_EVALUATION_HEADER = "PRUEBA DE EVALUACIÓN FINAL\nDEL MÓDULO (teórico-práctica)"
FINAL_EVALUATION_ITEMS = [
    (
        "1. Parte Teórica:",
        " Prueba teórica sobre los contenidos del manual y los contenidos "
        "específicos dados en las sesiones de formación.",
    ),
    (
        "2. Parte Práctica:",
        " Supuesto práctico completo del módulo",
    ),
]
RECOVERY_EVALUATION_HEADER = (
    "PRUEBA DE RECUPERACIÓN DE\n"
    "EVALUACIÓN FINAL DEL MÓDULO\n"
    "(teórico-práctica)"
)
RECOVERY_EVALUATION_INTRO = (
    "Re-evaluación de la Parte Teórica (POF) y\n"
    "la Parte Práctica (PPF)."
)
RECOVERY_EVALUATION_ITEMS = [
    (
        "1. Parte Teórica:",
        " Prueba teórica sobre los contenidos del manual y los contenidos "
        "específicos dados en las sesiones de formación.",
    ),
    (
        "2. Parte Práctica:",
        " Supuesto práctico completo del módulo",
    ),
]


def build_anexo_v_filename(module_code, certificate_code):
    return f"anexoV_{module_code}_{certificate_code}.docx"


def block_label(block):
    code = getattr(block, "code", "")
    name = getattr(block, "name", "")

    if code or name:
        return f"{code}: {name}".strip(": ")

    return module_identifier_without_hours(block)


def clean_evaluation_space_name(space):
    text = str(space or "").strip()
    text = re.sub(r"\s+de\s+\d+(?:,\d+)?\s*m2\b.*$", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s+\d+(?:,\d+)?\s*m2\b.*$", "", text, flags=re.IGNORECASE)
    return text.strip()


def evaluation_spaces_text(spaces):
    names = []
    seen = set()

    for space in spaces or []:
        name = clean_evaluation_space_name(space)
        key = name.lower()

        if name and key not in seen:
            names.append(name)
            seen.add(key)

    return "\n\n".join(names)


def evaluation_space_names(spaces):
    text = evaluation_spaces_text(spaces)
    return [name for name in text.split("\n\n") if name]


def set_cell_runs_color(cell, color):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = color


def set_top_cell_text(cell, text, align=WD_ALIGN_PARAGRAPH.LEFT):
    set_cell_text(
        cell,
        text,
        size=ANEXO_IV_FONT_SIZE,
        align=align,
        vertical_alignment=WD_ALIGN_VERTICAL.TOP,
    )


def set_activity_cell_text(cell, event):
    if event["type"] != "activity" or not event["label"]:
        set_top_cell_text(cell, event["label"])
        return

    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    title_paragraph = cell.paragraphs[0]
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_paragraph.paragraph_format.space_before = Pt(0)
    title_paragraph.paragraph_format.space_after = Pt(0)
    title_run = title_paragraph.add_run(event["label"])
    title_run.bold = True
    title_run.font.size = Pt(ANEXO_IV_FONT_SIZE)

    description_paragraph = cell.add_paragraph()
    description_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    description_paragraph.paragraph_format.space_before = Pt(0)
    description_paragraph.paragraph_format.space_after = Pt(0)
    description_run = description_paragraph.add_run(EVALUABLE_ACTIVITY_DESCRIPTION)
    description_run.font.size = Pt(ANEXO_IV_FONT_SIZE)
    description_run.font.color.rgb = RGBColor(192, 0, 0)


def set_red_cell_text(cell, text, align=WD_ALIGN_PARAGRAPH.CENTER):
    set_cell_text(cell, text, size=ANEXO_IV_FONT_SIZE, align=align)
    set_cell_runs_color(cell, RGBColor(192, 0, 0))


def set_duration_cell_text(cell, event, align=WD_ALIGN_PARAGRAPH.CENTER):
    set_cell_text(
        cell,
        duration_text(event["type"], event["session"]),
        size=ANEXO_IV_FONT_SIZE,
        align=align,
    )

    if event["type"] == "activity":
        set_cell_runs_color(cell, RGBColor(192, 0, 0))


def set_spaces_cell_text(cell, space_names, align=WD_ALIGN_PARAGRAPH.CENTER):
    text = "\n\n".join(space_names)

    if len(space_names) > 1:
        set_red_cell_text(cell, text, align=align)
    else:
        set_cell_text(cell, text, size=ANEXO_IV_FONT_SIZE, align=align)


def add_event_cells(cells, event, space_names=None):
    session = event["session"]
    space_names = space_names or []

    values = ["", event["label"], "", "", evaluation_date_text(session)]

    for index, value in enumerate(values):
        align = WD_ALIGN_PARAGRAPH.CENTER if index in (2, 3, 4) else WD_ALIGN_PARAGRAPH.LEFT

        if index == 2:
            set_spaces_cell_text(cells[index], space_names, align=align)
        elif index == 3:
            set_duration_cell_text(cells[index], event, align=align)
        elif index == 1:
            set_activity_cell_text(cells[index], event)
        elif index == 4:
            set_cell_text(cells[index], value, size=ANEXO_IV_FONT_SIZE, align=align)
        else:
            set_top_cell_text(cells[index], value, align=align)


def set_red_content_run(run, bold=False):
    run.bold = bold
    run.font.size = Pt(ANEXO_IV_FONT_SIZE)
    run.font.color.rgb = RGBColor(192, 0, 0)


def set_prefixed_red_items(cell, items, intro=None):
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    if intro:
        intro_paragraph = cell.paragraphs[0]
        intro_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        intro_paragraph.paragraph_format.space_before = Pt(0)
        intro_paragraph.paragraph_format.space_after = Pt(0)
        intro_run = intro_paragraph.add_run(intro)
        set_red_content_run(intro_run)

    for index, (prefix, text) in enumerate(items):
        paragraph = cell.paragraphs[0] if index == 0 and not intro else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

        prefix_run = paragraph.add_run(prefix)
        set_red_content_run(prefix_run, bold=True)

        text_run = paragraph.add_run(text)
        set_red_content_run(text_run)


def set_final_evaluation_content(cell):
    set_prefixed_red_items(cell, FINAL_EVALUATION_ITEMS)


def set_recovery_evaluation_content(cell):
    set_prefixed_red_items(
        cell,
        RECOVERY_EVALUATION_ITEMS,
        intro=RECOVERY_EVALUATION_INTRO,
    )


def prepare_row_cells(cells, widths):
    for index, cell in enumerate(cells):
        cell.width = widths[index]


def add_structured_evaluation_rows(
    table,
    event,
    widths,
    space_names,
    header_text,
    content_writer,
    dates_header="Fechas de\nevaluación²",
):
    header_cells = table.add_row().cells
    prepare_row_cells(header_cells, widths)
    set_top_cell_text(header_cells[0], "")

    for cell in header_cells[1:]:
        set_cell_shading(cell, ANEXO_IV_TABLE_HEADER_FILL)

    set_cell_text(header_cells[1], header_text, size=ANEXO_IV_FONT_SIZE)
    set_cell_text(header_cells[2], "Espacios", size=ANEXO_IV_FONT_SIZE)
    set_cell_text(header_cells[3], "Duración", size=ANEXO_IV_FONT_SIZE)
    set_cell_text(header_cells[4], dates_header, size=ANEXO_IV_FONT_SIZE)

    content_cells = table.add_row().cells
    prepare_row_cells(content_cells, widths)
    set_top_cell_text(content_cells[0], "")
    content_writer(content_cells[1])
    set_spaces_cell_text(content_cells[2], space_names)
    set_duration_cell_text(content_cells[3], event)
    set_cell_text(
        content_cells[4],
        evaluation_date_text(event["session"]),
        size=ANEXO_IV_FONT_SIZE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )


def add_final_event_rows(table, event, widths, space_names):
    add_structured_evaluation_rows(
        table,
        event,
        widths,
        space_names,
        FINAL_EVALUATION_HEADER,
        set_final_evaluation_content,
    )


def add_recovery_event_rows(table, event, widths, space_names):
    add_structured_evaluation_rows(
        table,
        event,
        widths,
        space_names,
        RECOVERY_EVALUATION_HEADER,
        set_recovery_evaluation_content,
        dates_header="Fechas de\nevaluación²",
    )


def add_anexo_v_table(doc, module, schedule=None, spaces=None):
    add_horizontal_line(doc)

    table = doc.add_table(rows=2, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    set_table_width_percent(table, ANEXO_IV_TABLE_WIDTH_PERCENT)

    widths = [Cm(3.1), Cm(6.1), Cm(2.4), Cm(1.8), Cm(2.4)]

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = widths[index]
            set_cell_shading(cell, ANEXO_IV_TABLE_HEADER_FILL)

    evaluation_header = table.cell(0, 2).merge(table.cell(0, 4))

    set_cell_text(table.cell(0, 0), "MÓDULO\nPROFESIONAL", size=ANEXO_IV_FONT_SIZE)
    set_cell_text(table.cell(1, 0), "BLOQUES\nFORMATIVOS", size=ANEXO_IV_FONT_SIZE)
    set_cell_text(table.cell(0, 1), "DURANTE EL PROCESO DE\nAPRENDIZAJE", size=ANEXO_IV_FONT_SIZE)
    set_cell_text(table.cell(1, 1), "ACTIVIDADES E INSTRUMENTOS\nDE EVALUACIÓN¹", size=ANEXO_IV_FONT_SIZE)
    set_cell_text(
        evaluation_header,
        "Realización de la evaluación",
        size=ANEXO_IV_FONT_SIZE,
    )

    subheaders = ["Espacios", "Duración", "Fechas de\nevaluación²"]

    for offset, header in enumerate(subheaders, start=2):
        cell = table.rows[1].cells[offset]
        set_cell_shading(cell, ANEXO_IV_TABLE_HEADER_FILL)
        set_cell_text(cell, header, size=ANEXO_IV_FONT_SIZE)
        cell.width = widths[offset]

    events_by_block = build_evaluation_events(module, schedule)
    space_names = evaluation_space_names(spaces)

    for block in evaluation_blocks(module):
        code = block_code(block)
        events = events_by_block.get(code) or [{
            "type": "activity",
            "label": "",
            "session": {"hours": 0, "session_hours": 0, "date": None, "session_number": ""},
        }]
        first_row_index = len(table.rows)

        for event in events:
            if event["type"] == "final":
                add_final_event_rows(table, event, widths, space_names)
                continue

            if event["type"] == "recovery":
                add_recovery_event_rows(table, event, widths, space_names)
                continue

            cells = table.add_row().cells
            prepare_row_cells(cells, widths)
            add_event_cells(cells, event, space_names)

        last_row_index = len(table.rows) - 1

        if last_row_index > first_row_index:
            block_cell = table.cell(first_row_index, 0).merge(table.cell(last_row_index, 0))
        else:
            block_cell = table.cell(first_row_index, 0)

        set_top_cell_text(block_cell, block_label(block))

    apply_vertical_borders(table, LIGHT_BORDER)
    return table


def add_anexo_v_notes(doc):
    notes = [
        (
            "1 Identificar las actividades e instrumentos de evaluación (E.; E.; etc.) "
            "indicando una denominación sintética de los mismos (supuestos prácticos, "
            "simulaciones, pruebas objetivas y/o pruebas de respuesta abierta)."
        ),
        (
            "2 Las fechas de evaluación estarán actualizadas en el momento en el que "
            "se efectúe la comunicación de inicio de las acciones formativas a la "
            "administración competente."
        ),
    ]

    doc.add_paragraph("")

    for note in notes:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(note)
        run.font.size = Pt(ANEXO_IV_FONT_SIZE)


def create_anexo_v_docx(
    data,
    module,
    duration_text,
    output_path,
    schedule=None,
    spaces=None,
    add_header_footer=None,
    teacher_name=DEFAULT_TEACHER_NAME,
    training_center=None,
):
    doc = Document()
    configure_page(doc.sections[0])

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(ANEXO_IV_FONT_SIZE)

    add_module_header(
        doc,
        data,
        module,
        duration_text,
        schedule,
        annex_label="ANEXO V",
        document_title="Planificación de la evaluación del aprendizaje",
        module_section_title="PLANIFICACIÓN DE LA EVALUACIÓN DEL APRENDIZAJE",
        training_center=training_center,
    )
    add_anexo_v_table(doc, module, schedule, spaces)
    add_anexo_v_notes(doc)

    if add_header_footer is None:
        from source.docx_utils import add_header_footer

    add_header_footer(doc, teacher_name)

    doc.save(output_path)
