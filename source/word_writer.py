import re

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from source.anexo_iii_writer import add_anexo_iii
from source.docx_utils import add_header_footer, add_horizontal_rule, safe_text


def add_paragraph_with_m2_superscript(doc, text):
    p = doc.add_paragraph()

    parts = re.split(r"(m2)", safe_text(text))

    for part in parts:
        if part == "m2":
            p.add_run("m")
            r = p.add_run("2")
            r.font.superscript = True
        else:
            p.add_run(part)

    return p


def add_separator_line(doc):
    add_horizontal_rule(doc, color="808080")


def add_bold_prefix_paragraph(doc, text, left_indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(left_indent)

    text = safe_text(text)
    match = re.match(r"^(C\d+:|CE\d+\.\d+)\s*(.*)", text)

    if match:
        p.add_run(match.group(1) + " ").bold = True
        p.add_run(match.group(2))
    else:
        p.add_run(text)

    return p


def add_criteria_block(doc, criteria):
    if not criteria:
        return

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.add_run("Capacidades y criterios de evaluación:").bold = True

    for criterion in criteria:
        criterion_text = criterion.text

        if criterion_text:
            add_bold_prefix_paragraph(doc, criterion_text, left_indent=18)

        for subcriterion in criterion.subcriteria:
            subcriterion_text = subcriterion.text

            if subcriterion_text:
                add_bold_prefix_paragraph(doc, subcriterion_text, left_indent=36)

            for bullet in subcriterion.bullets:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Pt(54)
                p.add_run(safe_text(bullet))


def add_content_bullet(doc, bullet, level=0):
    styles = ["List Bullet", "List Bullet 2", "List Bullet 3"]
    style = styles[min(level, len(styles) - 1)]

    p = doc.add_paragraph(style=style)
    p.paragraph_format.left_indent = Pt(54 + (level * 18))
    p.add_run(safe_text(bullet.text))

    for child in bullet.children:
        add_content_bullet(doc, child, level + 1)


def add_contents_block(doc, contents):
    if not contents:
        return

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.add_run("Contenidos").bold = True

    for content in contents:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)

        title = safe_text(content.title)
        match = re.match(r"^(\d+\.)\s*(.*)", title)
        if match:
            p.add_run(match.group(1) + " ").bold = True
            p.add_run(match.group(2)).bold = True
        else:
            p.add_run(title).bold = True

        for bullet in content.bullets:
            add_content_bullet(doc, bullet)


def create_info_docx(
    data,
    modules,
    spaces,
    equipment_groups,
    duration_text,
    training_modules,
    output_path,
    teacher_name="Docente"
):
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = safe_text(f"{data.codigo} - {data.nombre}".upper())

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Información obtenida del BOE. Por favor, revise que los datos son correctos.")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(255, 0, 0)

    doc.add_paragraph("")

    def bold_line(label, value):
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        p.add_run(safe_text(value))

    bold_line("Nombre del Certificado: ", data.nombre)
    bold_line("Código: ", data.codigo)
    bold_line("Familia profesional: ", data.familia)
    bold_line("Nivel: ", data.nivel)
    bold_line("Duración del certificado: ", duration_text)

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.add_run("Relación de módulos formativos y de unidades formativas:").bold = True

    letters = "abcdefghijklmnopqrstuvwxyz"

    for i, module in enumerate(modules, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.add_run(f"{i}. {module.text}")

        for j, uf in enumerate(module.ufs):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(54)
            p.add_run(f"{letters[j]}. {uf}")

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.add_run("Espacios, instalaciones y equipamiento:").bold = True

    p = doc.add_paragraph()
    p.add_run("Espacio formativo").bold = True

    for space in spaces:
        add_paragraph_with_m2_superscript(doc, space)

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.add_run("Equipamiento").bold = True

    for group in equipment_groups:
        p = doc.add_paragraph()
        p.add_run(group.name).bold = True

        for item in group.items:
            doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph("")
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PROGRAMACIÓN DIDÁCTICA DEL MÓDULO PROFESIONAL")
    r.bold = True
    r.font.size = Pt(14)

    for index, training_module in enumerate(training_modules):
        doc.add_paragraph("")

        bold_line("Identificación del módulo profesional: ", training_module.identifier)
        bold_line("Horas: ", f"{training_module.hours}h")
        bold_line("Objetivo general del módulo: ", training_module.objective)

        if training_module.ufs:
            for uf in training_module.ufs:
                p = doc.add_paragraph()
                p.add_run(f"Unidad formativa {uf.number}: ").bold = True
                p.add_run(f"{uf.code} {uf.name} ({uf.hours} horas)")

                add_criteria_block(doc, uf.criteria)
                add_contents_block(doc, uf.contents)
        else:
            add_criteria_block(doc, training_module.criteria)
            add_contents_block(doc, training_module.contents)

        if index < len(training_modules) - 1:
            add_separator_line(doc)

    add_header_footer(doc, teacher_name)

    doc.save(output_path)


def create_anexo_iii_docx(
    data,
    modules,
    duration_text,
    output_path,
    schedule=None,
    teacher_name="Docente"
):
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    add_anexo_iii(doc, data, modules, duration_text, schedule, new_page=False)
    add_header_footer(doc, teacher_name)

    doc.save(output_path)
