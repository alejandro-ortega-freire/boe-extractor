import re
import sys
import xml.etree.ElementTree as ET
from pathlib import Path
from zipfile import ZipFile

from docx import Document


TRUTH_FILES = {
    "MF1442_3": Path(r"D:\Programacion_Didactica_Modalidad_Presencial MF1442_3.odt"),
    "MF1443_3": Path(r"D:\Programacion_Didactica_Modalidad_Presencial MF1443_3.docx"),
    "MF1444_3": Path(r"D:\Programacion_Didactica_Modalidad_Presencial MF1444_3.docx"),
    "MF1445_3": Path(r"D:\Programacion_Didactica_Modalidad_Presencial MF1445_3.docx"),
    "MF1446_3": Path(r"D:\Programacion_Didactica_Modalidad_Presencial MF1446_3.docx"),
}


def clean(text):
    return " ".join(str(text or "").split())


def criterion_codes(text):
    return [f"C{match}" for match in re.findall(r"\b(?:RA|C)\s*(\d+)\b", text or "")]


def content_numbers(text):
    numbers = {
        int(match)
        for match in re.findall(r"(?:^|/|\s)(\d+)\s*\.", text or "")
    }
    numbers.update(
        int(match)
        for match in re.findall(r"\bUA\s*(\d+)\b", text or "", flags=re.IGNORECASE)
    )
    return numbers


def docx_rows(path):
    doc = Document(path)

    if not doc.tables:
        return []

    rows = []
    for row in doc.tables[0].rows:
        cells = []
        for cell in row.cells:
            text = " / ".join(
                paragraph.text.strip()
                for paragraph in cell.paragraphs
                if paragraph.text.strip()
            )
            cells.append(clean(text))
        rows.append(cells)

    return rows


def odt_rows(path):
    namespaces = {
        "table": "urn:oasis:names:tc:opendocument:xmlns:table:1.0",
        "text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0",
    }

    with ZipFile(path) as archive:
        root = ET.fromstring(archive.read("content.xml"))

    tables = root.findall(".//table:table", namespaces)

    if not tables:
        return []

    rows = []
    for row in tables[0].findall("table:table-row", namespaces):
        cells = []
        for cell in row.findall("table:table-cell", namespaces):
            texts = []
            for paragraph in cell.findall(".//text:p", namespaces):
                text = "".join(paragraph.itertext()).strip()

                if text:
                    texts.append(text)

            cells.append(clean(" / ".join(texts)))

        if any(cells):
            rows.append(cells)

    return rows


def table_rows(path):
    path = Path(path)

    if path.suffix.lower() == ".odt":
        return odt_rows(path)

    return docx_rows(path)


def mapping_from_anexo(path, fill_empty_with_previous=True):
    entries = []
    previous_numbers = set()

    for row in table_rows(path)[1:]:
        first_cell = row[0] if len(row) > 0 else ""
        second_cell = row[1] if len(row) > 1 else ""
        codes = criterion_codes(first_cell)
        numbers = content_numbers(second_cell)

        if not codes:
            continue

        if not numbers and fill_empty_with_previous:
            numbers = set(previous_numbers)

        if numbers:
            previous_numbers = set(numbers)

        for code in codes:
            entries.append({
                "criterion": code,
                "numbers": set(numbers),
            })

    return entries


def generated_file_for_module(output_dir, certificate_code, module_code):
    return (
        Path(output_dir)
        / certificate_code
        / "Anexos IV"
        / f"anexoIV_{module_code}_{certificate_code}.docx"
    )


def compare_mappings(expected, actual):
    details = []
    scores = []
    expected = list(expected)
    actual = list(actual)

    while len(expected) < len(actual):
        previous_numbers = expected[-1]["numbers"] if expected else set()
        previous_code = expected[-1]["criterion"] if expected else f"C{len(expected) + 1}"
        match = re.search(r"\d+", previous_code)
        next_code = f"C{int(match.group(0)) + 1}" if match else f"C{len(expected) + 1}"
        expected.append({
            "criterion": next_code,
            "numbers": set(previous_numbers),
        })

    max_length = max(len(expected), len(actual))

    for index in range(max_length):
        expected_entry = expected[index] if index < len(expected) else {"criterion": "", "numbers": set()}
        actual_entry = actual[index] if index < len(actual) else {"criterion": "", "numbers": set()}
        expected_numbers = expected_entry["numbers"]
        actual_numbers = actual_entry["numbers"]

        if not expected_numbers and not actual_numbers:
            score = 1.0
        elif not expected_numbers or not actual_numbers:
            score = 0.0
        else:
            score = len(expected_numbers & actual_numbers) / len(expected_numbers | actual_numbers)

        scores.append(score)
        details.append({
            "criterion": expected_entry["criterion"] or actual_entry["criterion"],
            "expected": sorted(expected_numbers),
            "actual": sorted(actual_numbers),
            "score": score,
        })

    average = sum(scores) / len(scores) if scores else 1.0
    return average, details


def compare_ssce0110(output_dir="output", certificate_code="SSCE0110"):
    results = []

    for module_code, truth_path in TRUTH_FILES.items():
        generated_path = generated_file_for_module(output_dir, certificate_code, module_code)

        if not truth_path.exists() or not generated_path.exists():
            results.append({
                "module": module_code,
                "score": 0.0,
                "missing": str(generated_path if not generated_path.exists() else truth_path),
                "details": [],
            })
            continue

        expected = mapping_from_anexo(truth_path)
        actual = mapping_from_anexo(generated_path, fill_empty_with_previous=False)
        score, details = compare_mappings(expected, actual)
        results.append({
            "module": module_code,
            "score": score,
            "details": details,
        })

    global_score = (
        sum(result["score"] for result in results) / len(results)
        if results
        else 0.0
    )
    return global_score, results


def print_report(output_dir="output", certificate_code="SSCE0110"):
    global_score, results = compare_ssce0110(output_dir, certificate_code)
    print(f"Similitud global: {global_score * 100:.2f}%")

    for result in results:
        print(f"\n{result['module']}: {result['score'] * 100:.2f}%")

        if result.get("missing"):
            print(f"  Falta archivo: {result['missing']}")
            continue

        for detail in result["details"]:
            if detail["score"] < 1.0:
                print(
                    "  "
                    f"{detail['criterion']}: esperado {detail['expected']} "
                    f"generado {detail['actual']} "
                    f"({detail['score'] * 100:.2f}%)"
                )


if __name__ == "__main__":
    output_dir = sys.argv[1] if len(sys.argv) > 1 else "output"
    certificate_code = sys.argv[2] if len(sys.argv) > 2 else "SSCE0110"
    print_report(output_dir, certificate_code)
