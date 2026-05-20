import re

from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor
from source.docx_styles import (
    ANEXO_FONT_SIZE,
    ANEXO_III_DATES_COLUMN_WIDTH_CM,
    ANEXO_III_HEADER_FILL,
    ANEXO_III_HEADER_ROW_HEIGHT_CM,
    ANEXO_III_HOURS_COLUMN_WIDTH_CM,
    ANEXO_III_TABLE_WIDTH_PERCENT,
)
from source.schedule import code_from_text, format_date_range, format_holiday_note
from source.settings import (
    ACTION_CODE,
    PLACEHOLDER_ADDRESS,
    PLACEHOLDER_CENTER,
    PLACEHOLDER_DATES,
    PLACEHOLDER_LOCALITY,
    PROVINCE,
)
from source.table_styles import (
    set_cell_shading,
    set_cell_text as set_table_cell_text,
    set_table_width_percent,
)


HOURS_COLUMN_WIDTH = Cm(ANEXO_III_HOURS_COLUMN_WIDTH_CM)
DATES_COLUMN_WIDTH = Cm(ANEXO_III_DATES_COLUMN_WIDTH_CM)


def is_practice_module(module):
    return code_from_text(module.text).startswith("MP")


def certificate_modules(modules):
    return [module for module in modules if not is_practice_module(module)]


def practice_modules(modules):
    return [module for module in modules if is_practice_module(module)]


def hours_from_text(text):
    found = re.findall(r"\((\d+)\s*horas?\)", text, flags=re.IGNORECASE)
    return found[-1] if found else ""


def title_without_hours(text):
    return re.sub(r"\s*\(\d+\s*horas?\)\s*\.?\s*$", "", text, flags=re.IGNORECASE).strip()


def duration_for_anexo(duration_text):
    match = re.match(r"^(\d+)h\s*\+\s*(\d+)h\s*FEM$", duration_text, flags=re.IGNORECASE)

    if match:
        mf_hours = int(match.group(1))
        mp_hours = int(match.group(2))
        return f"{mf_hours} + {mp_hours} FEM"

    match = re.match(r"^(\d+)h$", duration_text, flags=re.IGNORECASE)

    if match:
        return f"{match.group(1)} HORAS"

    return duration_text.upper()


def set_cell_text(cell, text, bold=False, size=ANEXO_FONT_SIZE, align=WD_ALIGN_PARAGRAPH.LEFT, vertical_padding=True):
    set_table_cell_text(
        cell,
        text,
        bold=bold,
        size=size,
        align=align,
        vertical_padding=vertical_padding,
    )


def set_header_cell(cell, text):
    set_cell_shading(cell, ANEXO_III_HEADER_FILL)
    set_cell_text(cell, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, vertical_padding=False)

    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)


def add_label_value(paragraph, label, value):
    paragraph.paragraph_format.space_after = Pt(0)
    label_run = paragraph.add_run(label)
    label_run.bold = True
    label_run.font.size = Pt(ANEXO_FONT_SIZE)
    value_run = paragraph.add_run(value)
    value_run.font.size = Pt(ANEXO_FONT_SIZE)


def add_centered_heading(doc, text, bold=True):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(ANEXO_FONT_SIZE)
    return paragraph


def add_tabbed_label_line(doc, parts):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)

    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(4.6))
    tab_stops.add_tab_stop(Inches(8.0))

    for index, (label, value) in enumerate(parts):
        if index:
            paragraph.add_run("\t")

        add_label_value(paragraph, label, value)

    return paragraph


def schedule_date_range(schedule):
    if not schedule:
        return PLACEHOLDER_DATES

    start = schedule.get("start_date")
    end = schedule.get("end_date")

    if not start or not end:
        return PLACEHOLDER_DATES

    return format_date_range(start, end)


def add_anexo_header(doc, data, duration_text, schedule=None):
    add_centered_heading(doc, "ANEXO III")
    add_centered_heading(doc, "Planificación didáctica")
    add_centered_heading(doc, "(Modalidad presencial)")
    doc.add_paragraph("")

    certificate = (
        f"{ACTION_CODE} {data.codigo} "
        f"{data.nombre.upper()}"
    ).strip()

    add_tabbed_label_line(doc, [
        ("CERTIFICADO DE PROFESIONALIDAD: ", certificate),
    ])
    add_tabbed_label_line(doc, [
        ("DURACIÓN DEL CERTIFICADO: ", duration_for_anexo(duration_text)),
        ("FECHAS DE IMPARTICIÓN: ", schedule_date_range(schedule)),
    ])
    add_tabbed_label_line(doc, [
        ("CENTRO DE FORMACIÓN: ", PLACEHOLDER_CENTER),
    ])
    add_tabbed_label_line(doc, [
        ("DIRECCIÓN: ", PLACEHOLDER_ADDRESS),
        ("LOCALIDAD: ", PLACEHOLDER_LOCALITY),
        ("PROVINCIA: ", PROVINCE),
    ])


def scheduled_text(schedule, text):
    if not schedule:
        return ""

    code = code_from_text(text)
    return schedule.get("dates_by_code", {}).get(code, {}).get("text", "")


def module_rows(module, schedule=None):
    module_text = title_without_hours(module.text)
    module_hours = hours_from_text(module.text)
    ufs = module.ufs

    if not ufs:
        return [{
            "module": module_text,
            "module_hours": module_hours,
            "uf": "",
            "uf_hours": "",
            "dates": scheduled_text(schedule, module.text),
        }]

    rows = []

    for uf in ufs:
        rows.append({
            "module": module_text,
            "module_hours": module_hours,
            "uf": title_without_hours(uf),
            "uf_hours": hours_from_text(uf),
            "dates": scheduled_text(schedule, uf),
        })

    return rows


def add_planning_table(doc, modules, schedule=None):
    doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("PLANIFICACIÓN DIDÁCTICA DEL CURSO COMPLETO")
    run.bold = True
    run.font.size = Pt(ANEXO_FONT_SIZE)

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    set_table_width_percent(table, ANEXO_III_TABLE_WIDTH_PERCENT)

    headers = [
        "MÓDULOS DEL CERTIFICADO",
        "HORAS\nDEL\nMÓDULO",
        "UNIDADES FORMATIVAS\n(UF)",
        "HORAS\n(UF)",
        "FECHAS DE IMPARTICIÓN",
    ]

    widths = [Inches(3.2), HOURS_COLUMN_WIDTH, Inches(3.6), HOURS_COLUMN_WIDTH, DATES_COLUMN_WIDTH]

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_header_cell(cell, header)
        cell.width = widths[index]

    table.rows[0].height = Cm(ANEXO_III_HEADER_ROW_HEIGHT_CM)

    for module in certificate_modules(modules):
        rows = module_rows(module, schedule)
        first_row_index = len(table.rows)

        for row_data in rows:
            cells = table.add_row().cells
            is_first_module_row = len(table.rows) - 1 == first_row_index
            values = [
                row_data["module"] if is_first_module_row else "",
                row_data["module_hours"] if is_first_module_row else "",
                row_data["uf"],
                row_data["uf_hours"],
                row_data["dates"],
            ]

            for index, value in enumerate(values):
                set_cell_text(
                    cells[index],
                    value,
                    size=ANEXO_FONT_SIZE,
                    align=WD_ALIGN_PARAGRAPH.CENTER if index in (1, 3, 4) else WD_ALIGN_PARAGRAPH.LEFT
                )
                cells[index].width = widths[index]

        if len(rows) > 1:
            last_row_index = len(table.rows) - 1
            module_cell = table.cell(first_row_index, 0).merge(table.cell(last_row_index, 0))
            hours_cell = table.cell(first_row_index, 1).merge(table.cell(last_row_index, 1))
            set_cell_text(module_cell, rows[0]["module"], size=ANEXO_FONT_SIZE)
            set_cell_text(hours_cell, rows[0]["module_hours"], size=ANEXO_FONT_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER)

    return table


def add_practice_table(doc, modules, schedule=None):
    practices = practice_modules(modules)

    if not practices:
        return None

    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    set_table_width_percent(table, ANEXO_III_TABLE_WIDTH_PERCENT)

    widths = [Inches(4.6), HOURS_COLUMN_WIDTH, Inches(5.0)]
    headers = [
        "Módulo de formación práctica\nen centros de trabajo",
        "HORAS\nDEL\nMÓDULO",
        "FECHAS DE REALIZACIÓN",
    ]

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_header_cell(cell, header)
        cell.width = widths[index]

        if index == 0:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.italic = True

    table.rows[0].height = Cm(ANEXO_III_HEADER_ROW_HEIGHT_CM)

    for module in practices:
        cells = table.add_row().cells
        values = [
            title_without_hours(module.text),
            hours_from_text(module.text),
            scheduled_text(schedule, module.text),
        ]

        for index, value in enumerate(values):
            set_cell_text(
                cells[index],
                value,
                size=ANEXO_FONT_SIZE,
                align=WD_ALIGN_PARAGRAPH.CENTER if index in (1, 2) else WD_ALIGN_PARAGRAPH.LEFT
            )
            cells[index].width = widths[index]

    doc.add_paragraph("")
    return table


def add_holiday_note(doc, schedule):
    if not schedule:
        return

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    run = paragraph.add_run(format_holiday_note(schedule))
    run.font.size = Pt(ANEXO_FONT_SIZE)


def configure_anexo_section(section):
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)


def add_anexo_iii(doc, data, modules, duration_text, schedule=None, new_page=True):
    section = doc.add_section(WD_SECTION.NEW_PAGE) if new_page else doc.sections[0]
    configure_anexo_section(section)

    add_anexo_header(doc, data, duration_text, schedule)
    add_planning_table(doc, modules, schedule)
    add_practice_table(doc, modules, schedule)
    add_holiday_note(doc, schedule)
