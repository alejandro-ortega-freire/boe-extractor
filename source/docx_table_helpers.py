from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from source.docx_styles import ANEXO_FONT_SIZE


BLACK_BORDER = "000000"
RED = RGBColor(255, 0, 0)
WHITE = RGBColor(255, 255, 255)


def add_title(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(ANEXO_FONT_SIZE)
    return paragraph


def set_cell_borders(cell, color=BLACK_BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")

    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{border_name}"))

        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)

        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def apply_table_borders(table, color=BLACK_BORDER):
    table_pr = table._tbl.tblPr
    borders = table_pr.find(qn("w:tblBorders"))

    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_pr.append(borders)

    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{border_name}"))

        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)

        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)

    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell, color)


def set_cell_runs_color(cell, color):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = color
