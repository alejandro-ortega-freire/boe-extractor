from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt

from source.anexo_iii_writer import add_anexo_header, certificate_modules, configure_anexo_section
from source.docx_styles import ANEXO_FONT_SIZE, ANEXO_III_HEADER_FILL, ANEXO_III_TABLE_WIDTH_PERCENT
from source.docx_table_helpers import RED, WHITE, add_title, apply_table_borders, set_cell_runs_color
from source.schedule import code_from_text
from source.settings import DEFAULT_STUDENT_COUNT, DEFAULT_TEACHER_NAME
from source.table_styles import set_cell_shading, set_cell_text, set_exact_row_height, set_table_width_percent


def module_code(module):
    return code_from_text(module_text(module))


def module_text(module):
    return getattr(module, "text", "") or getattr(module, "identifier", "")


def item_code(item):
    if isinstance(item, str):
        return code_from_text(item)

    return getattr(item, "code", "") or module_code(item)


def module_header_text(module, number):
    return f"MF {number}\n({module_code(module)})"


def uf_header_text(uf, number):
    return f"UF {number}\n({item_code(uf)})"


def practice_header_text(module):
    return f"MP\n({module_code(module)})"


def practice_module(modules):
    for module in modules:
        if code_from_text(module.text).startswith("MP"):
            return module

    return None


def acta_columns(modules):
    columns = [
        {"type": "index", "label": "Nº", "width": Inches(0.45)},
        {"type": "dni", "label": "DNI/NIE", "width": Inches(1.55)},
        {"type": "name", "label": "APELLIDOS/NOMBRE", "width": Inches(2.2)},
    ]
    module_number = 1

    for module in certificate_modules(modules):
        ufs = list(module.ufs or [])

        if ufs:
            for uf_index, _uf in enumerate(ufs, start=1):
                columns.append({
                    "type": "uf",
                    "module": module,
                    "module_number": module_number,
                    "label": uf_header_text(_uf, uf_index),
                    "value": "Apto /No apto",
                    "width": Inches(0.85),
                })

            columns.append({
                "type": "module_final",
                "module": module,
                "module_number": module_number,
                "label": "CALIFICACIÓN\nFINAL",
                "value": "Apto (suficiente)/\nApto (notable)/\nApto (sobresaliente)\n/No apto",
                "width": Inches(1.1),
            })
        else:
            columns.append({
                "type": "module_no_uf",
                "module": module,
                "module_number": module_number,
                "value": "Apto (suficiente)/\nApto (notable)/\nApto (sobresaliente)\n/No apto",
                "width": Inches(1.15),
            })

        module_number += 1

    practice = practice_module(modules)

    if practice:
        columns.append({
            "type": "practice",
            "label": practice_header_text(practice),
            "value": "Apto /No apto\n/Exento",
            "width": Inches(1.0),
        })

    columns.extend([
        {
            "type": "certificate_proposal",
            "label": "PROPUESTA\nCERTIFICADO",
            "value": "SI/NO",
            "width": Inches(1.2),
        },
        {
            "type": "partial_proposal",
            "label": "PROPUESTA\nACREDITACIÓN\nPARCIAL",
            "value": "SI/NO",
            "width": Inches(1.45),
        },
    ])

    return columns


def set_header_text(cell, text, align=WD_ALIGN_PARAGRAPH.CENTER):
    set_cell_shading(cell, ANEXO_III_HEADER_FILL)
    set_cell_text(
        cell,
        text,
        bold=True,
        size=ANEXO_FONT_SIZE,
        align=align,
        vertical_alignment=WD_ALIGN_VERTICAL.CENTER,
    )

    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = WHITE


def set_content_cell(cell, text, color=None, bold=True):
    set_cell_text(
        cell,
        text,
        bold=bold,
        size=ANEXO_FONT_SIZE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        vertical_alignment=WD_ALIGN_VERTICAL.CENTER,
    )

    if color is not None:
        set_cell_runs_color(cell, color)


def merge_vertical(table, start_row, end_row, col):
    return table.cell(start_row, col).merge(table.cell(end_row, col))


def add_acta_header(table, columns):
    for index, column in enumerate(columns):
        table.rows[0].cells[index].width = column["width"]
        table.rows[1].cells[index].width = column["width"]

    for index in range(3):
        cell = merge_vertical(table, 0, 1, index)
        set_header_text(cell, columns[index]["label"])

    index = 3

    while index < len(columns):
        column = columns[index]

        if column["type"] in ("uf", "module_final"):
            module = column["module"]
            module_number = column["module_number"]
            start = index

            while (
                index < len(columns)
                and columns[index].get("module") is module
                and columns[index]["type"] in ("uf", "module_final")
            ):
                index += 1

            header = table.cell(0, start).merge(table.cell(0, index - 1))
            set_header_text(header, module_header_text(module, module_number))

            for sub_index in range(start, index):
                set_header_text(table.cell(1, sub_index), columns[sub_index]["label"])

            continue

        if column["type"] == "module_no_uf":
            cell = merge_vertical(table, 0, 1, index)
            set_header_text(cell, module_header_text(column["module"], column["module_number"]))
        else:
            cell = merge_vertical(table, 0, 1, index)
            set_header_text(cell, column["label"])

        index += 1


def student_name_placeholder(student_number):
    if student_number == 1:
        return "Apellido 1 Apellido 2, Nombre"

    return ""


def dni_placeholder(student_number):
    if student_number == 1:
        return "00000000-L"

    return ""


def add_student_rows(table, columns, student_count):
    for student_number in range(1, student_count + 1):
        row = table.add_row()
        cells = row.cells

        if student_number >= 2:
            set_exact_row_height(row, Cm(1.4))

        for index, column in enumerate(columns):
            cells[index].width = column["width"]

            if column["type"] == "index":
                set_content_cell(cells[index], str(student_number), color=RED)
            elif column["type"] == "dni":
                set_content_cell(cells[index], dni_placeholder(student_number) if student_number == 1 else "", color=RED)
            elif column["type"] == "name":
                set_content_cell(
                    cells[index],
                    student_name_placeholder(student_number) if student_number == 1 else "",
                    color=RED,
                )
            else:
                set_content_cell(cells[index], column.get("value", "") if student_number == 1 else "", color=RED)


def add_anexo_vii_table(doc, modules, student_count=DEFAULT_STUDENT_COUNT):
    student_count = student_count or DEFAULT_STUDENT_COUNT
    add_title(doc, "ACTA DE EVALUACIÓN")

    columns = acta_columns(modules)
    table = doc.add_table(rows=2, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    set_table_width_percent(table, ANEXO_III_TABLE_WIDTH_PERCENT)

    add_acta_header(table, columns)
    add_student_rows(table, columns, student_count)
    apply_table_borders(table)
    return table


def create_anexo_vii_docx(
    data,
    modules,
    duration_text,
    output_path,
    schedule=None,
    teacher_name=DEFAULT_TEACHER_NAME,
    training_center=None,
    student_count=DEFAULT_STUDENT_COUNT,
):
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(ANEXO_FONT_SIZE)

    configure_anexo_section(doc.sections[0])
    add_anexo_header(
        doc,
        data,
        duration_text,
        schedule,
        annex_label="ANEXO VII",
        document_title="Acta de Evaluación",
        training_center=training_center,
    )
    add_anexo_vii_table(doc, modules, student_count=student_count)

    from source.docx_utils import add_header_footer

    add_header_footer(doc, teacher_name)
    doc.save(output_path)
