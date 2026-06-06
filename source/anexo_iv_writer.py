from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor
import re

from source.anexo_iii_writer import (
    duration_for_anexo,
    schedule_date_range,
    training_center_data,
)
from source.content_assignment import assign_contents_to_criteria
from source.docx_styles import (
    ANEXO_IV_FONT_SIZE,
    ANEXO_IV_MAIN_HEADER_ROW_HEIGHT_CM,
    ANEXO_IV_TABLE_WIDTH_PERCENT,
    ANEXO_IV_TABLE_HEADER_FILL,
    ANEXO_IV_UF_ROW_MIN_HEIGHT_CM,
    LIGHT_BORDER,
    SUGGESTION_COLOR,
    WHITE_FILL,
)
from source.docx_utils import add_horizontal_rule
from source.evaluation_plan import build_evaluation_events
from source.models import Criterion
from source.schedule import code_from_text, format_date_range
from source.settings import (
    DEFAULT_TEACHER_NAME,
)
from source.table_styles import (
    apply_vertical_borders,
    set_exact_row_height,
    set_cell_shading,
    set_cell_text as set_table_cell_text,
    set_minimum_row_height,
    set_table_width_percent,
)


UF_ROW_MIN_HEIGHT = Cm(ANEXO_IV_UF_ROW_MIN_HEIGHT_CM)
MAIN_HEADER_ROW_HEIGHT = Cm(ANEXO_IV_MAIN_HEADER_ROW_HEIGHT_CM)


def configure_page(section):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)


def apply_light_vertical_borders(table):
    apply_vertical_borders(table, LIGHT_BORDER)


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    set_table_cell_text(
        cell,
        text,
        bold=bold,
        size=ANEXO_IV_FONT_SIZE,
        align=align,
        vertical_alignment=WD_ALIGN_VERTICAL.CENTER,
    )


def clear_cell(cell):
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    return paragraph


def add_prefixed_text(paragraph, text, italic=False):
    text = str(text or "").strip()

    if not text:
        return

    parts = text.split(" ", 1)
    prefix = parts[0]
    rest = parts[1] if len(parts) > 1 else ""

    prefix_run = paragraph.add_run(prefix)
    prefix_run.bold = True
    prefix_run.italic = italic
    prefix_run.font.size = Pt(ANEXO_IV_FONT_SIZE)

    if rest:
        text_run = paragraph.add_run(" " + rest)
        text_run.italic = italic
        text_run.font.size = Pt(ANEXO_IV_FONT_SIZE)


def set_criterion_cell_text(cell, criterion, include_subcriteria=False):
    paragraph = clear_cell(cell)
    add_prefixed_text(paragraph, criterion.text)

    if not include_subcriteria:
        return

    if criterion.subcriteria:
        spacer = cell.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(0)

    for subcriterion in criterion.subcriteria:
        paragraph = cell.add_paragraph()
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        add_prefixed_text(paragraph, subcriterion.text, italic=True)

        for bullet in subcriterion.bullets:
            add_hyphen_item(cell, bullet, italic=True)


def add_bold_cell_paragraph(cell, text, space_before=0):
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(ANEXO_IV_FONT_SIZE)
    return paragraph


def add_cell_text_paragraph(cell, text):
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text or ""))
    run.font.size = Pt(ANEXO_IV_FONT_SIZE)
    return paragraph


def add_hyphen_item(cell, text, level=0, italic=False, color=None):
    paragraph = cell.add_paragraph()
    indent = 7 + (level * 8)
    paragraph.paragraph_format.left_indent = Pt(indent)
    paragraph.paragraph_format.first_line_indent = Pt(-7)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(f"- {text}")
    run.font.size = Pt(ANEXO_IV_FONT_SIZE)
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return paragraph


def set_spaces_equipment_cell_text(cell, spaces=None, equipment_groups=None):
    spaces = [space for space in (spaces or []) if space]
    equipment_items = []

    for group in equipment_groups or []:
        equipment_items.extend(item for item in group.items if item)

    clear_cell(cell)
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    cell.paragraphs[0].add_run("Espacio formativo").bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(ANEXO_IV_FONT_SIZE)

    if len(spaces) == 1:
        add_cell_text_paragraph(cell, spaces[0])
    else:
        for space in spaces:
            add_hyphen_item(cell, space)

    add_bold_cell_paragraph(cell, "Equipamiento", space_before=8)

    for item in equipment_items:
        add_hyphen_item(cell, item)


def add_content_bullets(cell, bullets, level=0, color=None):
    for bullet in bullets or []:
        text = bullet.text

        if text:
            add_hyphen_item(cell, text, level, color=color)

        add_content_bullets(cell, bullet.children, level + 1, color=color)


def set_contents_cell_text(cell, contents=None, suggested=False):
    contents = contents or []
    clear_cell(cell)
    color = SUGGESTION_COLOR if suggested else None

    if not contents:
        return

    for index, content in enumerate(contents):
        title = content.title

        if title:
            paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(6 if index else 0)
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(title)
            run.bold = True
            run.font.size = Pt(ANEXO_IV_FONT_SIZE)
            if color:
                run.font.color.rgb = color

        add_content_bullets(cell, content.bullets, color=color)


def strategy_suggestion_lines(module, schedule):
    if not schedule:
        return []

    events_by_block = build_evaluation_events(module, schedule)
    events = [
        event
        for block_events in events_by_block.values()
        for event in block_events
    ]
    events.sort(key=lambda event: event["session"]["session_number"])

    lines = [
        "Sugerencia:",
        "SESIÓN 1 ⭢ Presentación del MF",
    ]

    for event in events:
        session_number = event["session"].get(
            "module_session_number",
            event["session"].get("session_number"),
        )

        if event["type"] == "activity":
            match = re.match(r"^(E\d+)", event["label"])
            label = match.group(1) if match else event["label"]
        elif event["type"] == "final":
            label = "Prueba Final"
        elif event["type"] == "recovery":
            label = "Recuperación"
        else:
            label = event["label"]

        lines.append(f"SESIÓN {session_number} ⭢ {label}")

    return lines


def add_strategy_suggestion(cell, lines):
    if not lines:
        return

    red = RGBColor(192, 0, 0)

    cell.add_paragraph("")
    cell.add_paragraph("")

    for index, line in enumerate(lines):
        paragraph = cell.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        run.font.size = Pt(ANEXO_IV_FONT_SIZE)
        run.font.color.rgb = red
        run.bold = index == 0
        run.italic = index > 0


def set_strategy_placeholder_cell_text(cell, suggestion_lines=None):
    titles = [
        "SESIÓN #",
        "ESTRATEGIA\nMETODOLÓGICA",
        "Inicio (1h)",
        "Desarrollo (4h)",
        "Final (1h)",
        "ACTIVIDAD DE\nAPRENDIZAJE",
        "RECURSOS\nDIDÁCTICOS",
    ]

    clear_cell(cell)

    for index, title in enumerate(titles):
        if index == 0:
            paragraph = cell.paragraphs[0]
        else:
            paragraph = cell.add_paragraph()

        paragraph.paragraph_format.space_before = Pt(8 if index else 0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(title)
        run.bold = True
        run.font.size = Pt(ANEXO_IV_FONT_SIZE)

        text_paragraph = cell.add_paragraph()
        text_paragraph.paragraph_format.space_before = Pt(0)
        text_paragraph.paragraph_format.space_after = Pt(0)
        text_run = text_paragraph.add_run("Lorem ipsum dolor met")
        text_run.font.size = Pt(ANEXO_IV_FONT_SIZE)
        text_run.font.color.rgb = RGBColor(192, 0, 0)

    add_strategy_suggestion(cell, suggestion_lines)


def add_horizontal_line(doc):
    add_horizontal_rule(doc, color="000000", space_before=14, space_after=10)


def add_heading(doc, text, size=12, space_after=6):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    return paragraph


def add_label_value(paragraph, label, value):
    label_run = paragraph.add_run(label)
    label_run.bold = True
    label_run.font.size = Pt(ANEXO_IV_FONT_SIZE)
    value_run = paragraph.add_run(str(value or ""))
    value_run.font.size = Pt(ANEXO_IV_FONT_SIZE)


def add_tabbed_line(doc, parts, tab_stops=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)

    tab_stops = tab_stops or [Inches(3.6)]
    for stop in tab_stops:
        paragraph.paragraph_format.tab_stops.add_tab_stop(stop)

    for index, (label, value) in enumerate(parts):
        if index:
            paragraph.add_run("\t")

        add_label_value(paragraph, label, value)

    return paragraph


def scheduled_text(schedule, code):
    if not schedule:
        return ""

    return schedule.get("dates_by_code", {}).get(code, {}).get("text", "")


def module_title(module):
    return module_identifier_without_hours(module).replace(":", "", 1).strip()


def module_identifier_without_hours(module):
    identifier = module.identifier
    return re.sub(r"\s*\(\d+\s*horas?\)\s*\.?\s*$", "", identifier, flags=re.IGNORECASE).strip()


def module_schedule_text(schedule, module):
    module_code = code_from_text(module.identifier)
    direct_dates = scheduled_text(schedule, module_code)

    if direct_dates:
        return direct_dates

    uf_dates = []

    for uf in module.ufs:
        scheduled = schedule.get("dates_by_code", {}).get(uf.code) if schedule else None

        if scheduled:
            uf_dates.append(scheduled)

    if not uf_dates:
        return ""

    first = uf_dates[0]
    last = uf_dates[-1]
    return format_date_range(first["start"], last["end"])


def build_module_filename(module_code, certificate_code):
    return f"anexoIV_{module_code}_{certificate_code}.docx"


def add_module_header(
    doc,
    data,
    module,
    duration_text,
    schedule,
    annex_label="ANEXO IV",
    document_title="Programación didáctica",
    module_section_title="PROGRAMACIÓN DIDÁCTICA DEL MÓDULO PROFESIONAL",
    training_center=None,
):
    add_heading(doc, f"{annex_label} - {module_title(module)}", size=12, space_after=8)
    add_heading(doc, document_title, size=12, space_after=4)
    add_heading(doc, "(Modalidad Presencial)", size=12, space_after=16)

    certificate = f"{data.codigo} {data.nombre.upper()}".strip()

    add_tabbed_line(doc, [("CERTIFICADO PROFESIONAL: ", certificate)])
    add_tabbed_line(doc, [("FAMILIA PROFESIONAL: ", data.familia)])
    add_tabbed_line(doc, [("NIVEL DE CUALIFICACIÓN PROFESIONAL: ", data.nivel)])
    doc.add_paragraph("")

    add_tabbed_line(
        doc,
        [
            ("DURACIÓN DEL CERTIFICADO: ", duration_for_anexo(duration_text)),
            ("FECHAS DE IMPARTICIÓN: ", schedule_date_range(schedule)),
        ],
        tab_stops=[Inches(3.8)],
    )
    center = training_center_data(training_center)
    add_tabbed_line(doc, [("CENTRO DE FORMACIÓN: ", center["center"])])
    add_tabbed_line(doc, [("DIRECCIÓN: ", center["address"])])
    add_tabbed_line(
        doc,
        [
            ("LOCALIDAD: ", center["locality"]),
            ("PROVINCIA: ", center["province"]),
        ],
        tab_stops=[Inches(3.8)],
    )
    doc.add_paragraph("")

    add_heading(doc, module_section_title, size=11, space_after=8)

    module_dates = module_schedule_text(schedule, module)

    add_tabbed_line(doc, [("IDENTIFICACIÓN DEL MÓDULO PROFESIONAL: ", module_identifier_without_hours(module))])
    add_tabbed_line(
        doc,
        [
            ("HORAS: ", module.hours),
            ("FECHAS DE IMPARTICIÓN DEL MÓDULO: ", module_dates),
        ],
        tab_stops=[Inches(1.2)],
    )
    add_tabbed_line(doc, [("Nº DE CURSO: ", center["course_number"])])
    add_tabbed_line(doc, [("Objetivo general del módulo: ", module.objective)])


def add_anexo_iv_table(
    doc,
    module,
    schedule,
    copy_subcriteria=False,
    spaces=None,
    equipment_groups=None
):
    add_horizontal_line(doc)

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    set_table_width_percent(table, ANEXO_IV_TABLE_WIDTH_PERCENT)

    headers = [
        "Objetivos específicos\nLogro de los\nresultados de\naprendizaje ¹",
        "Contenidos²",
        "Estrategias\nmetodológicas,\nactividades de\naprendizaje y\nrecursos didácticos³",
        "Espacios,\ninstalaciones y\nequipamiento⁴",
    ]
    widths = [Cm(4.0), Cm(4.0), Cm(4.0), Cm(4.0)]

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, ANEXO_IV_TABLE_HEADER_FILL)
        set_cell_text(cell, header)
        cell.width = widths[index]

    set_exact_row_height(table.rows[0], MAIN_HEADER_ROW_HEIGHT)

    ufs = module.ufs
    suggestion_lines = strategy_suggestion_lines(module, schedule)
    suggestion_pending = bool(suggestion_lines)

    if not ufs:
        criteria = module.criteria or [Criterion()]
        contents_by_criterion = assign_contents_to_criteria(criteria, module.contents)
        suggested_contents = len(criteria) > 1

        for criterion_index, criterion in enumerate(criteria):
            cells = table.add_row().cells

            for index, cell in enumerate(cells):
                cell.width = widths[index]

            set_criterion_cell_text(cells[0], criterion, copy_subcriteria)

            for cell in cells[1:2]:
                set_cell_text(cell, "")

            set_contents_cell_text(
                cells[1],
                contents_by_criterion[criterion_index],
                suggested=suggested_contents
            )
            set_strategy_placeholder_cell_text(cells[2], suggestion_lines if suggestion_pending else None)
            suggestion_pending = False
            set_spaces_equipment_cell_text(cells[3], spaces, equipment_groups)

        apply_light_vertical_borders(table)
        return table

    for uf in ufs:
        uf_code = uf.code
        uf_name = uf.name
        uf_dates = scheduled_text(schedule, uf_code)

        top_cells = table.add_row().cells
        bottom_cells = table.add_row().cells
        set_minimum_row_height(table.rows[-2], UF_ROW_MIN_HEIGHT)
        set_minimum_row_height(table.rows[-1], UF_ROW_MIN_HEIGHT)

        for row_cells in (top_cells, bottom_cells):
            for index, cell in enumerate(row_cells):
                set_cell_shading(cell, ANEXO_IV_TABLE_HEADER_FILL)
                cell.width = widths[index]

        uf_label_cell = top_cells[0].merge(bottom_cells[0])
        uf_name_cell = top_cells[1].merge(bottom_cells[1])

        set_cell_text(uf_label_cell, "Unidad Formativa\n(UF)")
        set_cell_text(uf_name_cell, f"{uf_code}: {uf_name}".strip(": "))
        set_cell_shading(uf_name_cell, WHITE_FILL)

        set_cell_text(top_cells[2], "Horas")
        set_cell_text(top_cells[3], uf.hours, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(top_cells[3], WHITE_FILL)

        set_cell_text(bottom_cells[2], "Fechas de impartición")
        set_cell_text(bottom_cells[3], uf_dates, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(bottom_cells[3], WHITE_FILL)

        criteria = uf.criteria or [Criterion()]
        contents_by_criterion = assign_contents_to_criteria(criteria, uf.contents)
        suggested_contents = len(criteria) > 1

        for criterion_index, criterion in enumerate(criteria):
            cells = table.add_row().cells
            for index, cell in enumerate(cells):
                cell.width = widths[index]

            set_criterion_cell_text(cells[0], criterion, copy_subcriteria)
            for cell in cells[1:2]:
                set_cell_text(cell, "")

            set_contents_cell_text(
                cells[1],
                contents_by_criterion[criterion_index],
                suggested=suggested_contents
            )
            set_strategy_placeholder_cell_text(cells[2], suggestion_lines if suggestion_pending else None)
            suggestion_pending = False
            set_spaces_equipment_cell_text(cells[3], spaces, equipment_groups)

    apply_light_vertical_borders(table)
    return table


def add_anexo_iv_notes(doc):
    notes = [
        (
            "1  Incluir los Resultados de Aprendizaje (RA) tal y como se describen "
            "en el certificado profesional. En los certificados profesionales "
            "regulados en el plan antiguo los RA se corresponden con las Capacidades."
        ),
        (
            "2  Introducir los contenidos que se contemplan en el certificado, "
            "asignándolos a los RA correspondientes y secuenciándolos pedagógicamente."
        ),
        (
            "3  Especificar las diferentes acciones de enseñanza-aprendizaje que han "
            "de realizar los formadores y/o el alumnado para el logro de los resultados "
            "de aprendizaje, indicando los métodos didácticos a utilizar y los recursos "
            "didácticos asociados. Se incluyen también en este apartado las actividades "
            "de aprendizaje a realizar por los alumnos."
        ),
        (
            "4  Indicar los que corresponden exclusivamente a ese módulo, considerando "
            "lo establecido en la normativa correspondiente."
        ),
    ]

    doc.add_paragraph("")

    for note in notes:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(note)
        run.font.size = Pt(ANEXO_IV_FONT_SIZE)


def create_anexo_iv_docx(
    data,
    module,
    duration_text,
    output_path,
    schedule=None,
    add_header_footer=None,
    copy_subcriteria=False,
    spaces=None,
    equipment_groups=None,
    teacher_name=DEFAULT_TEACHER_NAME,
    training_center=None,
):
    doc = Document()
    configure_page(doc.sections[0])

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(ANEXO_IV_FONT_SIZE)

    add_module_header(doc, data, module, duration_text, schedule, training_center=training_center)
    add_anexo_iv_table(doc, module, schedule, copy_subcriteria, spaces, equipment_groups)
    add_anexo_iv_notes(doc)

    if add_header_footer is None:
        from source.docx_utils import add_header_footer

    add_header_footer(doc, teacher_name)

    doc.save(output_path)
