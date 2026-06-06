from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from source.anexo_iii_writer import (
    add_anexo_header,
    configure_anexo_section,
    set_header_cell,
    title_without_hours,
)
from source.evaluation_plan import (
    activity_numbers_by_block,
    evaluable_activity_count,
)
from source.docx_styles import ANEXO_FONT_SIZE, ANEXO_III_TABLE_WIDTH_PERCENT
from source.docx_styles import ANEXO_III_HEADER_FILL
from source.schedule import code_from_text, parse_hours
from source.settings import DEFAULT_TEACHER_NAME
from source.table_styles import set_cell_shading, set_cell_text, set_exact_row_height, set_table_width_percent


EMPTY_EVALUATION_FILL = "F2F2F2"
BLACK_BORDER = "000000"


def add_title(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(ANEXO_FONT_SIZE)
    return paragraph


def set_red_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    set_cell_text(cell, text, bold=bold, size=ANEXO_FONT_SIZE, align=align)

    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 0, 0)


def set_cell_borders(cell, color=BLACK_BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")

    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{border_name}"))

        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)

        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def apply_table_borders(table, color=BLACK_BORDER):
    table_pr = table._tbl.tblPr
    borders = table_pr.find(qn("w:tblBorders"))

    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_pr.append(borders)

    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{border_name}"))

        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)

        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)

    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell, color)


def set_main_header_with_subtitle(cell, title, subtitle):
    set_cell_shading(cell, ANEXO_III_HEADER_FILL)
    set_header_cell(cell, "")
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_run = paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(ANEXO_FONT_SIZE)
    title_run.font.color.rgb = RGBColor(255, 255, 255)

    paragraph.add_run("\n")
    subtitle_run = paragraph.add_run(subtitle)
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(ANEXO_FONT_SIZE)
    subtitle_run.font.color.rgb = RGBColor(255, 255, 255)


def module_label(module_text):
    text = title_without_hours(module_text or "")
    return text


def uf_label(uf_text):
    text = title_without_hours(uf_text or "")
    return text


def certificate_modules(modules):
    return [
        module
        for module in modules
        if not code_from_text(module.text).startswith("MP")
    ]


def max_evaluable_count(modules):
    counts = [
        evaluable_activity_count(parse_hours(module.text))
        for module in certificate_modules(modules)
    ]
    return max(counts, default=1)


def set_row_widths(row, widths):
    for index, cell in enumerate(row.cells):
        cell.width = widths[index]


def set_score_cell(cell):
    set_red_text(cell, "Nota", align=WD_ALIGN_PARAGRAPH.CENTER)


def set_summary_score_cell(cell):
    set_red_text(cell, "Nota final", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def set_average_cell(cell, e_count):
    set_red_text(cell, f"ΣEvaluables / {e_count}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def set_evaluation_cells(cells, e_count, activity_numbers):
    e_start = 2
    average_col = e_start + e_count
    first_call_col = average_col + 1
    second_call_col = first_call_col + 2
    final_grade_col = second_call_col + 2

    for activity_number in range(1, e_count + 1):
        if activity_number in activity_numbers:
            set_score_cell(cells[e_start + activity_number - 1])
        else:
            set_cell_shading(cells[e_start + activity_number - 1], EMPTY_EVALUATION_FILL)

    for index in (first_call_col, first_call_col + 1, second_call_col, second_call_col + 1):
        set_score_cell(cells[index])

    for index in (second_call_col, second_call_col + 1):
        set_cell_text(cells[index], "", size=ANEXO_FONT_SIZE)
        set_cell_shading(cells[index], EMPTY_EVALUATION_FILL)

    set_red_text(cells[final_grade_col], "APTO (puntuación\nfinal)/NO APTO", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def set_final_summary_cells(row, e_count):
    cells = row.cells
    average_col = 2 + e_count
    first_call_col = average_col + 1
    second_call_col = first_call_col + 2

    first_summary = cells[first_call_col].merge(cells[first_call_col + 1])
    second_summary = cells[second_call_col].merge(cells[second_call_col + 1])
    set_summary_score_cell(first_summary)
    set_cell_text(second_summary, "", size=ANEXO_FONT_SIZE)
    set_cell_shading(second_summary, EMPTY_EVALUATION_FILL)


def add_module_rows(table, modules, widths, e_count, schedule):
    for module in certificate_modules(modules):
        ufs = module.ufs or []
        first_row_index = len(table.rows)
        activities_by_block = activity_numbers_by_block(module, schedule)
        module_evaluable_count = evaluable_activity_count(parse_hours(module.text))

        if not ufs:
            row = table.add_row()
            set_row_widths(row, widths)
            set_exact_row_height(row, Cm(1.4))
            set_evaluation_cells(
                row.cells,
                e_count,
                activities_by_block.get(code_from_text(module.text), set()),
            )
            summary_row = table.add_row()
            set_row_widths(summary_row, widths)
            set_final_summary_cells(summary_row, e_count)

            last_row_index = len(table.rows) - 1
            module_cell = table.cell(first_row_index, 0).merge(table.cell(last_row_index, 1))
            set_cell_text(module_cell, module_label(module.text), size=ANEXO_FONT_SIZE, align=WD_ALIGN_PARAGRAPH.LEFT)

            for index in range(2, 2 + e_count):
                table.cell(first_row_index, index).merge(table.cell(last_row_index, index))

            average_cell = table.cell(first_row_index, 2 + e_count).merge(table.cell(last_row_index, 2 + e_count))
            set_average_cell(average_cell, module_evaluable_count)
            final_grade_col = 2 + e_count + 5
            final_grade_cell = table.cell(first_row_index, final_grade_col).merge(table.cell(last_row_index, final_grade_col))
            set_red_text(final_grade_cell, "APTO (puntuación\nfinal)/NO APTO", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            continue

        for uf in ufs:
            row = table.add_row()
            set_row_widths(row, widths)
            set_cell_text(row.cells[1], uf_label(uf), size=ANEXO_FONT_SIZE, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_evaluation_cells(row.cells, e_count, activities_by_block.get(code_from_text(uf), set()))

        module_cell = table.cell(first_row_index, 0).merge(table.cell(len(table.rows) - 1, 0))
        set_cell_text(module_cell, module_label(module.text), size=ANEXO_FONT_SIZE, align=WD_ALIGN_PARAGRAPH.LEFT)
        average_cell = table.cell(first_row_index, 2 + e_count).merge(table.cell(len(table.rows) - 1, 2 + e_count))
        set_average_cell(average_cell, module_evaluable_count)
        final_grade_col = 2 + e_count + 5
        final_grade_cell = table.cell(first_row_index, final_grade_col).merge(table.cell(len(table.rows) - 1, final_grade_col))
        set_red_text(final_grade_cell, "APTO (puntuación\nfinal)/NO APTO", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_final_summary_cells(table.rows[-1], e_count)


def add_anexo_vi_table(doc, modules, schedule=None):
    add_title(doc, "INFORME DE EVALUACIÓN INDIVIDUALIZADO")

    e_count = max_evaluable_count(modules)
    column_count = e_count + 8

    table = doc.add_table(rows=4, cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    set_table_width_percent(table, ANEXO_III_TABLE_WIDTH_PERCENT)

    widths = [
        Inches(2.1),
        Inches(1.5),
        *[Inches(0.85) for _ in range(e_count)],
        Inches(1.25),
        Inches(0.75),
        Inches(0.75),
        Inches(0.75),
        Inches(0.75),
        Inches(1.55),
    ]

    for row in table.rows:
        set_row_widths(row, widths)

    name_label = table.cell(0, 0).merge(table.cell(0, 1))
    name_value = table.cell(0, 2).merge(table.cell(0, column_count - 1))
    set_cell_text(name_label, "Nombre y apellidos del alumno", bold=True, size=ANEXO_FONT_SIZE)
    set_red_text(name_value, "Nombre del alumno Apellido 1 Apellido 2")
    set_exact_row_height(table.rows[0], Cm(0.75))

    modules_header = table.cell(1, 0).merge(table.cell(3, 1))
    e_start = 2
    average_col = e_start + e_count
    first_call_col = average_col + 1
    second_call_col = first_call_col + 2
    final_grade_col = second_call_col + 2

    process_header = table.cell(1, e_start).merge(table.cell(1, average_col))
    final_header = table.cell(1, first_call_col).merge(table.cell(1, second_call_col + 1))
    final_grade_header = table.cell(1, final_grade_col).merge(table.cell(3, final_grade_col))

    set_header_cell(modules_header, "MÓDULOS FORMATIVOS")
    set_main_header_with_subtitle(
        process_header,
        "EVALUACIÓN DURANTE EL PROCESO DE APRENDIZAJE",
        "30% de la calificación final",
    )
    set_main_header_with_subtitle(
        final_header,
        "PRUEBA DE EVALUACIÓN FINAL\nDEL MÓDULO",
        "70% de la calificación final",
    )
    set_header_cell(final_grade_header, "CALIFICACIÓN FINAL DEL\nMÓDULO")

    results_header = table.cell(2, e_start).merge(table.cell(2, e_start + e_count - 1))
    set_header_cell(results_header, "Resultados en las actividades e instrumentos de evaluación")
    average_header = table.cell(2, average_col).merge(table.cell(3, average_col))
    set_header_cell(average_header, "Puntuación\nmedia")
    first_call_header = table.cell(2, first_call_col).merge(table.cell(2, first_call_col + 1))
    set_header_cell(first_call_header, "PRUEBA FINAL")
    second_call_header = table.cell(2, second_call_col).merge(table.cell(2, second_call_col + 1))
    set_header_cell(second_call_header, "RECUPERACIÓN")

    for offset in range(e_count):
        set_header_cell(table.cell(3, e_start + offset), f"E{offset + 1}")

    for index, label in {
        first_call_col: "POF\n40%",
        first_call_col + 1: "PPF\n60%",
        second_call_col: "POF\n40%",
        second_call_col + 1: "PPF\n60%",
    }.items():
        set_header_cell(table.cell(3, index), label)

    add_module_rows(table, modules, widths, e_count, schedule)
    apply_table_borders(table)
    return table


def create_anexo_vi_docx(
    data,
    modules,
    duration_text,
    output_path,
    schedule=None,
    teacher_name=DEFAULT_TEACHER_NAME,
    training_center=None,
):
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(ANEXO_FONT_SIZE)

    configure_anexo_section(doc.sections[0])
    add_anexo_header(
        doc,
        data,
        duration_text,
        schedule,
        annex_label="ANEXO VI",
        document_title="Informe de Evaluación Individualizado",
        training_center=training_center,
    )
    add_anexo_vi_table(doc, modules, schedule)

    from source.docx_utils import add_header_footer

    add_header_footer(doc, teacher_name)
    doc.save(output_path)
